import base64
import io
import re
import zipfile
from pathlib import Path
from typing import Any, Dict, List


class PdfConversionError(RuntimeError):
    def __init__(self, message: str, status_code: int = 500):
        super().__init__(message)
        self.status_code = status_code


def convert_pdf(pdf_bytes: bytes, original_filename: str, target_format: str) -> Dict[str, Any]:
    fmt = (target_format or "").strip().lower()
    if fmt not in {"md", "docx"}:
        raise PdfConversionError("仅支持转换为 md 或 docx", 400)
    if not pdf_bytes:
        raise PdfConversionError("PDF 文件不能为空", 400)
    if not pdf_bytes.lstrip().startswith(b"%PDF"):
        raise PdfConversionError("上传文件不是有效的 PDF", 400)

    base_name = _safe_stem(original_filename)
    if fmt == "docx":
        return _convert_to_docx(pdf_bytes, base_name)
    return _convert_to_markdown_zip(pdf_bytes, base_name)


def _convert_to_docx(pdf_bytes: bytes, base_name: str) -> Dict[str, Any]:
    try:
        import fitz
        from docx import Document
        from docx.shared import Inches
    except Exception as exc:
        raise PdfConversionError("PDF 转 DOCX 依赖未安装，请安装 PyMuPDF 和 python-docx", 500) from exc

    try:
        pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception as exc:
        raise PdfConversionError(f"PDF 解析失败：{exc}", 400) from exc

    document = Document()
    document.add_heading(base_name, level=0)
    image_assets: List[Dict[str, Any]] = []
    has_extractable_text = False
    page_count = pdf_document.page_count

    try:
        for page_index in range(page_count):
            page = pdf_document.load_page(page_index)
            page_number = page_index + 1
            document.add_heading(f"第 {page_number} 页", level=1)

            page_dict = page.get_text("dict")
            text_blocks = [block for block in page_dict.get("blocks", []) if block.get("type") == 0]
            image_blocks = [block for block in page_dict.get("blocks", []) if block.get("type") == 1]

            for block in sorted(text_blocks + image_blocks, key=_block_position):
                if block.get("type") == 0:
                    if _append_docx_text_block(document, block):
                        has_extractable_text = True
                    continue
                asset = _append_docx_image_block(document, block, page_number, len(image_assets) + 1, Inches)
                if asset:
                    image_assets.append(asset)

            if page_index < page_count - 1:
                document.add_page_break()
    finally:
        pdf_document.close()

    if not has_extractable_text:
        raise PdfConversionError("该 PDF 未检测到可提取文本；当前仅支持原生文字 PDF，不支持扫描件 OCR", 422)

    buffer = io.BytesIO()
    document.save(buffer)
    output_bytes = buffer.getvalue()
    return {
        "format": "docx",
        "outputType": "file",
        "downloadType": "file",
        "fileName": f"{base_name}.docx",
        "mimeType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "contentBase64": base64.b64encode(output_bytes).decode("ascii"),
        "contentLength": len(output_bytes),
        "assets": image_assets,
        "imageCount": len(image_assets),
        "pageCount": page_count,
        "conversionMode": "pdf_to_docx_reflow",
    }


def _convert_to_markdown_zip(pdf_bytes: bytes, base_name: str) -> Dict[str, Any]:
    try:
        import fitz
    except Exception as exc:
        raise PdfConversionError("PDF 转 Markdown 依赖未安装，请安装 PyMuPDF", 500) from exc

    assets: List[Dict[str, Any]] = []
    markdown_lines: List[str] = [f"# {base_name}", ""]
    has_extractable_text = False

    try:
        document = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception as exc:
        raise PdfConversionError(f"PDF 解析失败：{exc}", 400) from exc

    try:
        for page_index in range(document.page_count):
            page = document.load_page(page_index)
            page_number = page_index + 1
            page_text = (page.get_text("text") or "").strip()
            page_images = []

            if page_text:
                has_extractable_text = True

            for image_index, image in enumerate(page.get_images(full=True), start=1):
                xref = image[0]
                try:
                    image_info = document.extract_image(xref)
                except Exception as exc:
                    raise PdfConversionError(f"第 {page_number} 页图片提取失败：{exc}", 500) from exc
                image_bytes = image_info.get("image") or b""
                if not image_bytes:
                    continue
                ext = _safe_extension(image_info.get("ext") or "png")
                image_name = f"page-{page_number}-image-{image_index}.{ext}"
                relative_path = f"assets/{image_name}"
                mime_type = _image_mime_type(ext)
                page_images.append(relative_path)
                assets.append({
                    "name": image_name,
                    "path": relative_path,
                    "type": "image",
                    "mimeType": mime_type,
                    "previewDataUrl": f"data:{mime_type};base64,{base64.b64encode(image_bytes).decode('ascii')}",
                    "page": page_number,
                    "size": len(image_bytes),
                    "_bytes": image_bytes,
                })

            markdown_lines.extend([f"## 第 {page_number} 页", ""])
            if page_text:
                markdown_lines.extend([page_text, ""])
            for relative_path in page_images:
                markdown_lines.extend([f"![{Path(relative_path).name}]({relative_path})", ""])
    finally:
        document.close()

    if not has_extractable_text:
        raise PdfConversionError("该 PDF 未检测到可提取文本；当前版本不做 OCR，请先对扫描件做 OCR 后再转换", 422)

    markdown = "\n".join(markdown_lines).strip() + "\n"
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(f"{base_name}.md", markdown)
        for asset in assets:
            archive.writestr(asset["path"], asset["_bytes"])

    public_assets = [
        {key: value for key, value in asset.items() if key != "_bytes"}
        for asset in assets
    ]
    output_bytes = zip_buffer.getvalue()
    return {
        "format": "md",
        "outputType": "markdown",
        "downloadType": "zip",
        "fileName": f"{base_name}-markdown.zip",
        "mimeType": "application/zip",
        "contentBase64": base64.b64encode(output_bytes).decode("ascii"),
        "contentLength": len(output_bytes),
        "preview": markdown,
        "assets": public_assets,
        "imageCount": len(public_assets),
    }


def _safe_stem(filename: str) -> str:
    stem = Path(filename or "converted").stem or "converted"
    safe = re.sub(r'[\\/:*?"<>|\x00-\x1f]+', "-", stem).strip(".- ")
    return safe or "converted"


def _block_position(block: Dict[str, Any]) -> tuple:
    bbox = block.get("bbox") or (0, 0, 0, 0)
    return (float(bbox[1] or 0), float(bbox[0] or 0))


def _block_text(block: Dict[str, Any]) -> str:
    lines: List[str] = []
    for line in block.get("lines", []):
        spans = line.get("spans", [])
        text = "".join(span.get("text", "") for span in spans).strip()
        if text:
            lines.append(text)
    return "\n".join(lines).strip()


def _append_docx_text_block(document: Any, block: Dict[str, Any]) -> bool:
    text = _block_text(block)
    if not text:
        return False
    for line in text.splitlines():
        if line.strip():
            document.add_paragraph(line.strip())
    return True


def _append_docx_image_block(
    document: Any,
    block: Dict[str, Any],
    page_number: int,
    image_index: int,
    inches_factory: Any,
) -> Dict[str, Any]:
    image_bytes = block.get("image") or b""
    if not image_bytes:
        return {}
    ext = _safe_extension(block.get("ext") or "png")
    name = f"page-{page_number}-image-{image_index}.{ext}"
    bbox = block.get("bbox") or (0, 0, 0, 0)
    width_points = max(float(bbox[2] or 0) - float(bbox[0] or 0), 72.0)
    width_inches = max(1.0, min(width_points / 72.0, 6.2))
    document.add_picture(io.BytesIO(image_bytes), width=inches_factory(width_inches))
    return {
        "name": name,
        "path": f"assets/{name}",
        "type": "image",
        "mimeType": _image_mime_type(ext),
        "page": page_number,
        "size": len(image_bytes),
    }


def _safe_extension(ext: str) -> str:
    safe = re.sub(r"[^A-Za-z0-9]+", "", ext.lower())
    return safe or "png"


def _image_mime_type(ext: str) -> str:
    return {
        "jpg": "image/jpeg",
        "jpeg": "image/jpeg",
        "png": "image/png",
        "gif": "image/gif",
        "webp": "image/webp",
        "bmp": "image/bmp",
        "tif": "image/tiff",
        "tiff": "image/tiff",
    }.get(ext, f"image/{ext}")
