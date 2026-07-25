import hashlib
import hmac
import json
import mimetypes
import os
import re
import secrets
import stat
import tempfile
import uuid
import zipfile
from dataclasses import dataclass, field
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any, BinaryIO, Dict, List, Mapping, Optional, Set
from xml.sax.saxutils import escape

from docx import Document

DEFAULT_EXPORT_TTL_HOURS = 168.0
DEFAULT_EXPORT_MAX_BYTES = 1024 * 1024 * 1024
DEFAULT_EXPORT_STAGING_GRACE_SECONDS = 300.0
EXPORT_URL_PATH = "/uploads/ai-exports"


def _resolve_export_root(environment: Optional[Mapping[str, str]] = None) -> Path:
    values = os.environ if environment is None else environment
    configured_root = str(values.get("AI_EXPORT_ROOT") or "").strip()
    deployment_environment = str(values.get("AI_ENV") or "").strip().lower()
    if deployment_environment in {"prod", "production"} and not configured_root:
        raise RuntimeError(
            "AI_EXPORT_ROOT must be explicitly configured to a persistent shared export store in production"
        )
    if configured_root:
        return Path(configured_root).expanduser().resolve()
    return (Path(__file__).resolve().parents[3] / "data" / "ai-exports").resolve()


def _positive_number_setting(name: str, default: float) -> float:
    raw_value = str(os.getenv(name) or "").strip()
    if not raw_value:
        return default
    try:
        value = float(raw_value)
    except ValueError as exc:
        raise RuntimeError(f"{name} must be a positive number") from exc
    if value <= 0:
        raise RuntimeError(f"{name} must be a positive number")
    return value


EXPORT_ROOT = _resolve_export_root()
EXPORT_TTL_HOURS = _positive_number_setting("AI_EXPORT_TTL_HOURS", DEFAULT_EXPORT_TTL_HOURS)
EXPORT_MAX_BYTES = int(_positive_number_setting("AI_EXPORT_MAX_BYTES", DEFAULT_EXPORT_MAX_BYTES))
EXPORT_STAGING_GRACE_SECONDS = _positive_number_setting(
    "AI_EXPORT_STAGING_GRACE_SECONDS",
    DEFAULT_EXPORT_STAGING_GRACE_SECONDS,
)

_STORAGE_KEY_PATTERN = re.compile(
    r"^(?P<uuid>[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12})"
    r"\.(?P<extension>[a-z0-9]{1,16})$"
)
_MANIFEST_SUFFIX = ".meta.json"
_MANIFEST_FIELDS = {
    "capabilityDigest",
    "sha256",
    "size",
    "mimeType",
    "createdAt",
    "expiresAt",
}
_MIME_TYPES = {
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "zip": "application/zip",
    "md": "text/markdown",
    "mmd": "text/plain",
}

EXPORTABLE_MARKDOWN_ANSWER_TYPES = {
    "markdown",
    "ppt_outline",
    "ppt_layout",
    "ppt_review",
    "tool_result",
}
EXPORTABLE_MARKDOWN_AGENTS = {
    "textbook_knowledge_agent",
    "meeting_controller_agent",
    "meeting_transcription_agent",
    "meeting_summary_agent",
    "meeting_member_analysis_agent",
    "meeting_resource_recommendation_agent",
    "meeting_voice_broadcast_agent",
    "ppt_outline_agent",
    "ppt_layout_agent",
    "ppt_review_agent",
}
EXPORTABLE_DIAGRAM_ANSWER_TYPES = {
    "mermaid_mindmap",
    "mermaid_flowchart",
    "mermaid_activity_flowchart",
    "mermaid_architecture",
}

GENERATED_EXPORT_TOOL_NAME = "generated_export_tools"
MARKDOWN_EXPORT_TOOL_NAME = "markdown_export_tool"
DOCX_EXPORT_TOOL_NAME = "docx_export_tool"
EXCEL_EXPORT_TOOL_NAME = "excel_export_tool"
ARCHIVE_EXPORT_TOOL_NAME = "content_archive_tool"
DIAGRAM_SOURCE_EXPORT_TOOL_NAME = "diagram_source_export_tool"
KNOWN_EXPORT_TOOL_NAMES = {
    GENERATED_EXPORT_TOOL_NAME,
    MARKDOWN_EXPORT_TOOL_NAME,
    DOCX_EXPORT_TOOL_NAME,
    EXCEL_EXPORT_TOOL_NAME,
    ARCHIVE_EXPORT_TOOL_NAME,
    DIAGRAM_SOURCE_EXPORT_TOOL_NAME,
}


@dataclass
class GeneratedExportResult:
    attachments: List[Dict[str, Any]] = field(default_factory=list)
    diagnostics: Dict[str, Any] = field(default_factory=dict)


@dataclass(frozen=True)
class GeneratedExportFile:
    stream: BinaryIO
    storage_key: str
    mime_type: str
    sha256: str
    size: int
    created_at: str
    expires_at: str


class GeneratedExportAccessError(Exception):
    def __init__(self, status_code: int, detail: str) -> None:
        super().__init__(detail)
        self.status_code = status_code
        self.detail = detail


def export_generated_answer(answer: str, answer_type: str, metadata: Optional[Dict[str, Any]] = None) -> GeneratedExportResult:
    cleanup_generated_exports()
    metadata = metadata or {}
    agent = str(metadata.get("executedAgent") or metadata.get("targetAgent") or metadata.get("agentName") or "").strip()
    normalized_type = str(answer_type or metadata.get("answerType") or "").strip()
    content = str(answer or "").strip()
    if not content:
        return GeneratedExportResult(diagnostics={"skipped": True, "reason": "empty_answer"})
    if not metadata.get("allowGeneratedExportTool") and (agent == "generated_export_tools" or normalized_type == "document_export"):
        return GeneratedExportResult(diagnostics={"skipped": True, "reason": "already_exported_by_generated_export_tools"})
    if not _is_export_tool_enabled(metadata, GENERATED_EXPORT_TOOL_NAME):
        return GeneratedExportResult(diagnostics={
            "skipped": True,
            "reason": "tool_disabled",
            "disabledTool": GENERATED_EXPORT_TOOL_NAME,
        })

    if normalized_type == "question_bank" or agent.startswith("textbook_question_"):
        payload = _parse_json_object(content)
        if not payload or not isinstance(payload.get("questions"), list):
            return GeneratedExportResult(diagnostics={"skipped": True, "reason": "invalid_question_bank_json"})
        return _finalize_export_batch(_export_question_bank(payload, metadata))

    if normalized_type in EXPORTABLE_DIAGRAM_ANSWER_TYPES or _extract_mermaid_code(content):
        return _finalize_export_batch(_export_diagram_source(content, metadata))

    if _should_export_markdown(normalized_type, agent, metadata, content):
        return _finalize_export_batch(_export_markdown_content(content, metadata))

    return GeneratedExportResult(diagnostics={"skipped": True, "reason": "not_exportable_answer_type"})


def _should_export_markdown(answer_type: str, agent: str, metadata: Dict[str, Any], content: str) -> bool:
    requested = str(metadata.get("requestedOutputType") or metadata.get("preferredOutputType") or "").strip().lower()
    if requested in {"document", "file", "docx", "word", "excel", "md", "markdown"}:
        return True
    if agent in EXPORTABLE_MARKDOWN_AGENTS:
        return True
    if answer_type in EXPORTABLE_MARKDOWN_ANSWER_TYPES and _looks_like_markdown(content):
        return True
    return False


def _is_export_tool_enabled(metadata: Dict[str, Any], tool_name: str) -> bool:
    toggles = metadata.get("toolToggles")
    if not isinstance(toggles, dict):
        return True
    if tool_name not in toggles:
        return True
    return _parse_enabled_value(toggles.get(tool_name))


def _parse_enabled_value(value: Any) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return value != 0
    text = str(value or "").strip().lower()
    return text not in {"0", "false", "off", "disabled", "no"}


def _disabled_export_tools(metadata: Dict[str, Any]) -> List[str]:
    toggles = metadata.get("toolToggles")
    if not isinstance(toggles, dict):
        return []
    return [
        name
        for name in sorted(KNOWN_EXPORT_TOOL_NAMES)
        if name in toggles and not _parse_enabled_value(toggles.get(name))
    ]


def _formats_from_attachments(attachments: List[Dict[str, Any]]) -> List[str]:
    formats: List[str] = []
    for attachment in attachments:
        ext = str(attachment.get("ext") or "").strip().lower()
        if ext and ext not in formats:
            formats.append(ext)
    return formats


def _no_enabled_export_format_result(content_kind: str, metadata: Dict[str, Any], extra: Optional[Dict[str, Any]] = None) -> GeneratedExportResult:
    diagnostics = {
        "skipped": True,
        "reason": "no_enabled_export_format",
        "contentKind": content_kind,
        "disabledTools": _disabled_export_tools(metadata),
    }
    if extra:
        diagnostics.update(extra)
    return GeneratedExportResult(diagnostics=diagnostics)


def _export_question_bank(payload: Dict[str, Any], metadata: Dict[str, Any]) -> GeneratedExportResult:
    title = _title_from_metadata(metadata, "题库导出")
    slug = _slugify(title or "question-bank")
    markdown = _question_bank_to_markdown(payload, title)
    rows = _question_bank_rows(payload)
    paths: List[Path] = []
    attachments: List[Dict[str, Any]] = []
    if _is_export_tool_enabled(metadata, MARKDOWN_EXPORT_TOOL_NAME):
        path = _write_text_file(slug, "md", markdown)
        paths.append(path)
        attachments.append(_attachment_for_file(path, MARKDOWN_EXPORT_TOOL_NAME, "Markdown"))
    if _is_export_tool_enabled(metadata, DOCX_EXPORT_TOOL_NAME):
        path = _write_question_bank_docx(slug, title, payload)
        paths.append(path)
        attachments.append(_attachment_for_file(path, DOCX_EXPORT_TOOL_NAME, "Word 文档"))
    if _is_export_tool_enabled(metadata, EXCEL_EXPORT_TOOL_NAME):
        path = _write_xlsx(slug, "题库", rows)
        paths.append(path)
        attachments.append(_attachment_for_file(path, EXCEL_EXPORT_TOOL_NAME, "Excel 表格"))
    if _is_export_tool_enabled(metadata, ARCHIVE_EXPORT_TOOL_NAME) and len(paths) >= 2:
        attachments.append(_attachment_for_file(_write_archive(slug, paths), ARCHIVE_EXPORT_TOOL_NAME, "打包文件"))
    if not attachments:
        return _no_enabled_export_format_result(
            "question_bank",
            metadata,
            {"questionCount": len(payload.get("questions") or [])},
        )
    return GeneratedExportResult(
        attachments=attachments,
        diagnostics={
            "skipped": False,
            "contentKind": "question_bank",
            "questionCount": len(payload.get("questions") or []),
            "producedFormats": _formats_from_attachments(attachments),
            "disabledTools": _disabled_export_tools(metadata),
        },
    )


def _export_markdown_content(content: str, metadata: Dict[str, Any]) -> GeneratedExportResult:
    title = _title_from_markdown(content) or _title_from_metadata(metadata, "知识整理")
    slug = _slugify(title or "knowledge")
    rows = _markdown_rows(content)
    paths: List[Path] = []
    attachments: List[Dict[str, Any]] = []
    if _is_export_tool_enabled(metadata, MARKDOWN_EXPORT_TOOL_NAME):
        path = _write_text_file(slug, "md", content)
        paths.append(path)
        attachments.append(_attachment_for_file(path, MARKDOWN_EXPORT_TOOL_NAME, "Markdown"))
    if _is_export_tool_enabled(metadata, DOCX_EXPORT_TOOL_NAME):
        path = _write_markdown_docx(slug, title, content)
        paths.append(path)
        attachments.append(_attachment_for_file(path, DOCX_EXPORT_TOOL_NAME, "Word 文档"))
    if rows and _is_export_tool_enabled(metadata, EXCEL_EXPORT_TOOL_NAME):
        path = _write_xlsx(slug, "知识清单", rows)
        paths.append(path)
        attachments.append(_attachment_for_file(path, EXCEL_EXPORT_TOOL_NAME, "Excel 表格"))
    if _is_export_tool_enabled(metadata, ARCHIVE_EXPORT_TOOL_NAME) and len(paths) >= 2:
        attachments.append(_attachment_for_file(_write_archive(slug, paths), ARCHIVE_EXPORT_TOOL_NAME, "打包文件"))
    if not attachments:
        return _no_enabled_export_format_result(
            "markdown_content",
            metadata,
            {"itemCount": max(len(rows) - 1, 0)},
        )
    return GeneratedExportResult(
        attachments=attachments,
        diagnostics={
            "skipped": False,
            "contentKind": "markdown_content",
            "itemCount": max(len(rows) - 1, 0),
            "producedFormats": _formats_from_attachments(attachments),
            "disabledTools": _disabled_export_tools(metadata),
        },
    )


def _export_diagram_source(content: str, metadata: Dict[str, Any]) -> GeneratedExportResult:
    title = _title_from_metadata(metadata, "图表源码")
    slug = _slugify(title or "diagram-source")
    mermaid_code = _extract_mermaid_code(content) or str(content or "").strip()
    markdown = f"# {title or '图表源码'}\n\n```mermaid\n{mermaid_code}\n```\n"
    paths: List[Path] = []
    attachments: List[Dict[str, Any]] = []
    if _is_export_tool_enabled(metadata, DIAGRAM_SOURCE_EXPORT_TOOL_NAME):
        path = _write_text_file(slug, "mmd", mermaid_code.strip() + "\n")
        paths.append(path)
        attachments.append(_attachment_for_file(path, DIAGRAM_SOURCE_EXPORT_TOOL_NAME, "Mermaid 源文件"))
    if _is_export_tool_enabled(metadata, MARKDOWN_EXPORT_TOOL_NAME):
        path = _write_text_file(f"{slug}-mermaid", "md", markdown)
        paths.append(path)
        attachments.append(_attachment_for_file(path, MARKDOWN_EXPORT_TOOL_NAME, "Markdown"))
    if _is_export_tool_enabled(metadata, ARCHIVE_EXPORT_TOOL_NAME) and len(paths) >= 2:
        attachments.append(_attachment_for_file(_write_archive(slug, paths), ARCHIVE_EXPORT_TOOL_NAME, "打包文件"))
    if not attachments:
        return _no_enabled_export_format_result("diagram_source", metadata)
    return GeneratedExportResult(
        attachments=attachments,
        diagnostics={
            "skipped": False,
            "contentKind": "diagram_source",
            "producedFormats": _formats_from_attachments(attachments),
            "disabledTools": _disabled_export_tools(metadata),
        },
    )


def _write_text_file(slug: str, ext: str, content: str) -> Path:
    path = _new_export_path(slug, ext)
    _atomic_write_payload(path, lambda temporary_path: temporary_path.write_text(content, encoding="utf-8"))
    return path


def _write_markdown_docx(slug: str, title: str, content: str) -> Path:
    path = _new_export_path(slug, "docx")
    doc = Document()
    doc.add_heading(title or "知识整理", level=1)
    in_code_block = False
    for raw_line in str(content or "").splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            paragraph = doc.add_paragraph()
            paragraph.add_run(line).font.name = "Consolas"
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            level = min(len(heading.group(1)) + 1, 4)
            doc.add_heading(_clean_inline_markdown(heading.group(2)), level=level)
            continue
        if re.match(r"^[-*]\s+", stripped):
            doc.add_paragraph(_clean_inline_markdown(re.sub(r"^[-*]\s+", "", stripped)), style="List Bullet")
            continue
        if re.match(r"^\d+[.)]\s+", stripped):
            doc.add_paragraph(_clean_inline_markdown(re.sub(r"^\d+[.)]\s+", "", stripped)), style="List Number")
            continue
        doc.add_paragraph(_clean_inline_markdown(stripped))
    _atomic_write_payload(path, doc.save)
    return path


def _write_question_bank_docx(slug: str, title: str, payload: Dict[str, Any]) -> Path:
    path = _new_export_path(slug, "docx")
    doc = Document()
    doc.add_heading(title or "题库导出", level=1)
    questions = payload.get("questions") if isinstance(payload.get("questions"), list) else []
    for index, question in enumerate(questions, start=1):
        if not isinstance(question, dict):
            continue
        qtype = str(question.get("type") or "").strip()
        stem = str(question.get("stem") or question.get("title") or "").strip()
        doc.add_heading(f"{index}. {stem or '未命名题目'}", level=2)
        doc.add_paragraph(f"题型：{qtype or '-'}")
        if question.get("score") is not None:
            doc.add_paragraph(f"分值：{question.get('score')}")
        if question.get("difficulty"):
            doc.add_paragraph(f"难度：{question.get('difficulty')}")
        options = _extract_options(question)
        if options:
            doc.add_paragraph("选项：")
            for option in options:
                doc.add_paragraph(f"{option.get('key') or ''}. {option.get('text') or ''}".strip(), style="List Bullet")
        answer = _format_answer(question.get("answer"))
        if answer:
            doc.add_paragraph(f"答案：{answer}")
        analysis = str(question.get("analysis") or question.get("explanation") or "").strip()
        if analysis:
            doc.add_paragraph(f"解析：{analysis}")
        knowledge = _join_values(question.get("knowledgePoints"))
        if knowledge:
            doc.add_paragraph(f"知识点：{knowledge}")
        tags = _join_values(question.get("tags"))
        if tags:
            doc.add_paragraph(f"标签：{tags}")
    _atomic_write_payload(path, doc.save)
    return path


def _write_xlsx(slug: str, sheet_name: str, rows: List[List[Any]]) -> Path:
    path = _new_export_path(slug, "xlsx")
    safe_rows = rows if rows else [["内容"]]
    sheet_xml = _sheet_xml(safe_rows)
    escaped_sheet_name = escape(_safe_sheet_name(sheet_name), {'"': "&quot;"})
    workbook_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <sheets><sheet name="{escaped_sheet_name}" sheetId="1" r:id="rId1"/></sheets>
</workbook>"""
    workbook_rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet1.xml"/>
  <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>
</Relationships>"""
    root_rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>
</Relationships>"""
    content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>
  <Override PartName="/xl/worksheets/sheet1.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>
  <Override PartName="/xl/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.styles+xml"/>
</Types>"""
    styles = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<styleSheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
  <fonts count="1"><font><sz val="11"/><name val="Calibri"/></font></fonts>
  <fills count="1"><fill><patternFill patternType="none"/></fill></fills>
  <borders count="1"><border><left/><right/><top/><bottom/><diagonal/></border></borders>
  <cellStyleXfs count="1"><xf numFmtId="0" fontId="0" fillId="0" borderId="0"/></cellStyleXfs>
  <cellXfs count="1"><xf numFmtId="0" fontId="0" fillId="0" borderId="0" xfId="0"/></cellXfs>
  <cellStyles count="1"><cellStyle name="Normal" xfId="0" builtinId="0"/></cellStyles>
</styleSheet>"""
    def write_workbook(temporary_path: Path) -> None:
        with zipfile.ZipFile(temporary_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
            archive.writestr("[Content_Types].xml", content_types)
            archive.writestr("_rels/.rels", root_rels)
            archive.writestr("xl/workbook.xml", workbook_xml)
            archive.writestr("xl/_rels/workbook.xml.rels", workbook_rels)
            archive.writestr("xl/worksheets/sheet1.xml", sheet_xml)
            archive.writestr("xl/styles.xml", styles)

    _atomic_write_payload(path, write_workbook)
    return path


def _write_archive(slug: str, paths: List[Path]) -> Path:
    path = _new_export_path(f"{slug}-bundle", "zip")
    def write_archive(temporary_path: Path) -> None:
        with zipfile.ZipFile(temporary_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
            for item in paths:
                if item and item.exists():
                    archive.write(item, arcname=item.name)

    _atomic_write_payload(path, write_archive)
    return path


def _sheet_xml(rows: List[List[Any]]) -> str:
    row_xml = []
    for row_index, row in enumerate(rows, start=1):
        cells = []
        for col_index, value in enumerate(row, start=1):
            cell_ref = f"{_column_name(col_index)}{row_index}"
            cells.append(f'<c r="{cell_ref}" t="inlineStr"><is><t>{escape(_cell_text(value))}</t></is></c>')
        row_xml.append(f'<row r="{row_index}">{"".join(cells)}</row>')
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
  <sheetViews><sheetView workbookViewId="0"/></sheetViews>
  <sheetData>{"".join(row_xml)}</sheetData>
</worksheet>"""


def _question_bank_to_markdown(payload: Dict[str, Any], title: str) -> str:
    lines = [f"# {title or '题库导出'}", ""]
    questions = payload.get("questions") if isinstance(payload.get("questions"), list) else []
    for index, question in enumerate(questions, start=1):
        if not isinstance(question, dict):
            continue
        lines.append(f"## {index}. {question.get('stem') or question.get('title') or '未命名题目'}")
        lines.append("")
        lines.append(f"- 题型：{question.get('type') or '-'}")
        if question.get("score") is not None:
            lines.append(f"- 分值：{question.get('score')}")
        if question.get("difficulty"):
            lines.append(f"- 难度：{question.get('difficulty')}")
        options = _extract_options(question)
        if options:
            lines.append("- 选项：")
            lines.extend(f"  - {option.get('key') or ''}. {option.get('text') or ''}".rstrip() for option in options)
        answer = _format_answer(question.get("answer"))
        if answer:
            lines.append(f"- 答案：{answer}")
        analysis = str(question.get("analysis") or question.get("explanation") or "").strip()
        if analysis:
            lines.append(f"- 解析：{analysis}")
        knowledge = _join_values(question.get("knowledgePoints"))
        if knowledge:
            lines.append(f"- 知识点：{knowledge}")
        tags = _join_values(question.get("tags"))
        if tags:
            lines.append(f"- 标签：{tags}")
        lines.append("")
    missing = payload.get("missingInfo")
    if isinstance(missing, list) and missing:
        lines.append("## 缺失信息")
        lines.extend(f"- {item}" for item in missing)
    return "\n".join(lines).strip() + "\n"


def _question_bank_rows(payload: Dict[str, Any]) -> List[List[Any]]:
    rows = [[
        "序号",
        "题型",
        "题干",
        "选项",
        "答案",
        "解析",
        "分值",
        "难度",
        "知识点",
        "标签",
    ]]
    questions = payload.get("questions") if isinstance(payload.get("questions"), list) else []
    for index, question in enumerate(questions, start=1):
        if not isinstance(question, dict):
            continue
        options = "；".join(f"{item.get('key') or ''}.{item.get('text') or ''}".strip(".") for item in _extract_options(question))
        rows.append([
            index,
            question.get("type") or "",
            question.get("stem") or question.get("title") or "",
            options,
            _format_answer(question.get("answer")),
            question.get("analysis") or question.get("explanation") or "",
            question.get("score") if question.get("score") is not None else "",
            question.get("difficulty") or "",
            _join_values(question.get("knowledgePoints")),
            _join_values(question.get("tags")),
        ])
    return rows


def _markdown_rows(content: str) -> List[List[Any]]:
    rows = [["序号", "类型", "层级", "内容"]]
    for line in str(content or "").splitlines():
        stripped = line.strip()
        if not stripped or stripped.startswith("```"):
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            rows.append([len(rows), "标题", len(heading.group(1)), _clean_inline_markdown(heading.group(2))])
            continue
        if re.match(r"^[-*]\s+", stripped):
            rows.append([len(rows), "要点", "", _clean_inline_markdown(re.sub(r"^[-*]\s+", "", stripped))])
            continue
        if re.match(r"^\d+[.)]\s+", stripped):
            rows.append([len(rows), "步骤", "", _clean_inline_markdown(re.sub(r"^\d+[.)]\s+", "", stripped))])
            continue
        if len(stripped) <= 240:
            rows.append([len(rows), "正文", "", _clean_inline_markdown(stripped)])
        if len(rows) >= 201:
            break
    return rows if len(rows) > 1 else []


def _attachment_for_file(path: Path, tool_name: str, format_label: str) -> Dict[str, Any]:
    ext = path.suffix.lower().lstrip(".")
    attachment_type = "docx" if ext == "docx" else "excel" if ext == "xlsx" else "file"
    export_metadata = _commit_export_manifest(path)
    return {
        "name": path.name,
        "fileName": path.name,
        "type": attachment_type,
        "ext": ext,
        "mimeType": export_metadata["mimeType"],
        "toolName": tool_name,
        "formatLabel": format_label,
        "source": "generated_content_export",
        "storageKey": path.name,
        "serverGenerated": True,
        "internalCapability": export_metadata["internalCapability"],
        "sha256": export_metadata["sha256"],
        "size": export_metadata["size"],
        "createdAt": export_metadata["createdAt"],
        "expiresAt": export_metadata["expiresAt"],
    }


def _new_export_path(slug: str, ext: str) -> Path:
    EXPORT_ROOT.mkdir(parents=True, exist_ok=True)
    del slug
    normalized_extension = re.sub(r"[^a-z0-9]", "", str(ext or "").lower())[:16]
    if not normalized_extension:
        raise ValueError("generated export extension is required")
    filename = f"{uuid.uuid4()}.{normalized_extension}"
    return EXPORT_ROOT / filename


def _atomic_write_payload(path: Path, writer: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    descriptor, temporary_name = tempfile.mkstemp(
        dir=path.parent,
        prefix=f".{path.name}.",
        suffix=".tmp",
    )
    os.close(descriptor)
    temporary_path = Path(temporary_name)
    try:
        writer(temporary_path)
        with temporary_path.open("rb") as stream:
            _fsync_if_supported(stream)
        os.replace(temporary_path, path)
    finally:
        _safe_unlink(temporary_path)


def _atomic_write_json(path: Path, value: Dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    descriptor, temporary_name = tempfile.mkstemp(
        dir=path.parent,
        prefix=f".{path.name}.",
        suffix=".tmp",
    )
    temporary_path = Path(temporary_name)
    try:
        with os.fdopen(descriptor, "w", encoding="utf-8") as stream:
            json.dump(value, stream, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
            stream.flush()
            _fsync_if_supported(stream)
        os.replace(temporary_path, path)
    except Exception:
        try:
            os.close(descriptor)
        except OSError:
            pass
        raise
    finally:
        _safe_unlink(temporary_path)


def _fsync_if_supported(stream: Any) -> None:
    try:
        fileno = stream.fileno()
    except (AttributeError, OSError, ValueError):
        return
    try:
        os.fsync(fileno)
    except OSError:
        # Windows can reject fsync on some temporary handles even after a successful write.
        # The export should still succeed once the payload has been flushed and replaced.
        pass


def _commit_export_manifest(path: Path) -> Dict[str, Any]:
    size = path.stat().st_size
    if size > EXPORT_MAX_BYTES:
        _safe_unlink(path)
        raise RuntimeError("generated export exceeds AI_EXPORT_MAX_BYTES")
    capability = secrets.token_urlsafe(32)
    created_at = datetime.now(timezone.utc)
    expires_at = created_at + timedelta(hours=EXPORT_TTL_HOURS)
    extension = path.suffix.lower().lstrip(".")
    mime_type = _MIME_TYPES.get(extension) or mimetypes.guess_type(path.name)[0] or "application/octet-stream"
    manifest = {
        "capabilityDigest": hashlib.sha256(capability.encode("utf-8")).hexdigest(),
        "sha256": _sha256_file(path),
        "size": size,
        "mimeType": mime_type,
        "createdAt": _format_utc(created_at),
        "expiresAt": _format_utc(expires_at),
    }
    _atomic_write_json(_manifest_path(path), manifest)
    return {**manifest, "internalCapability": capability}


def _finalize_export_batch(result: GeneratedExportResult) -> GeneratedExportResult:
    storage_keys = {
        str(attachment.get("storageKey") or "")
        for attachment in result.attachments
        if attachment.get("storageKey")
    }
    generated_size = sum(
        int(attachment.get("size") or 0)
        for attachment in result.attachments
        if attachment.get("storageKey") in storage_keys
    )
    if generated_size > EXPORT_MAX_BYTES:
        for storage_key in storage_keys:
            _delete_export_pair(EXPORT_ROOT, storage_key)
        raise RuntimeError("generated export batch exceeds AI_EXPORT_MAX_BYTES")
    cleanup_generated_exports(preserve_storage_keys=storage_keys)
    return result


def cleanup_generated_exports(
    *,
    root: Optional[Path] = None,
    now: Optional[datetime] = None,
    max_bytes: Optional[int] = None,
    preserve_storage_keys: Optional[Set[str]] = None,
    staging_grace_seconds: Optional[float] = None,
) -> None:
    export_root = Path(root or EXPORT_ROOT).resolve()
    export_root.mkdir(parents=True, exist_ok=True)
    current_time = _as_utc(now or datetime.now(timezone.utc))
    capacity = EXPORT_MAX_BYTES if max_bytes is None else max(0, int(max_bytes))
    staging_grace = (
        EXPORT_STAGING_GRACE_SECONDS
        if staging_grace_seconds is None
        else max(0.0, float(staging_grace_seconds))
    )
    preserved = set(preserve_storage_keys or set())
    payloads: Dict[str, Path] = {}
    manifests: Dict[str, Path] = {}

    try:
        entries = list(export_root.iterdir())
    except OSError:
        return
    for entry in entries:
        if entry.is_symlink():
            _safe_unlink(entry)
            continue
        if not entry.is_file():
            continue
        storage_key = _storage_key_from_manifest_name(entry.name)
        if storage_key:
            manifests[storage_key] = entry
        elif _is_valid_storage_key(entry.name):
            payloads[entry.name] = entry
        elif _is_older_than_grace(entry, current_time, staging_grace):
            _safe_unlink(entry)

    complete_pairs = set(payloads) & set(manifests)
    for storage_key in (set(payloads) | set(manifests)) - complete_pairs:
        payload_path = payloads.get(storage_key)
        manifest_path = manifests.get(storage_key)
        if payload_path is not None and _is_older_than_grace(payload_path, current_time, staging_grace):
            _safe_unlink(payload_path)
        if manifest_path is not None and _is_older_than_grace(manifest_path, current_time, staging_grace):
            _safe_unlink(manifest_path)

    retained = []
    for storage_key in complete_pairs:
        payload_path = payloads[storage_key]
        manifest_path = manifests[storage_key]
        try:
            manifest = _load_manifest(manifest_path)
            expires_at = _parse_utc(manifest["expiresAt"])
            created_at = _parse_utc(manifest["createdAt"])
        except (OSError, ValueError, TypeError, json.JSONDecodeError):
            try:
                created_at = datetime.fromtimestamp(payload_path.stat().st_mtime, tz=timezone.utc)
            except OSError:
                _delete_export_pair(export_root, storage_key)
                continue
            expires_at = None
        if expires_at is not None and expires_at <= current_time and storage_key not in preserved:
            _delete_export_pair(export_root, storage_key)
            continue
        try:
            size = payload_path.stat().st_size
        except OSError:
            _delete_export_pair(export_root, storage_key)
            continue
        retained.append((created_at, storage_key, size))

    total_size = sum(item[2] for item in retained)
    for _, storage_key, size in sorted(retained, key=lambda item: (item[0], item[1])):
        if total_size <= capacity:
            break
        if storage_key in preserved:
            continue
        _delete_export_pair(export_root, storage_key)
        total_size -= size


def open_generated_export(
    storage_key: str,
    capability: Optional[str],
    *,
    root: Optional[Path] = None,
    now: Optional[datetime] = None,
) -> GeneratedExportFile:
    export_root = Path(root or EXPORT_ROOT).resolve()
    normalized_key = str(storage_key or "")
    if not _is_valid_storage_key(normalized_key):
        raise GeneratedExportAccessError(404, "generated export not found")
    cleanup_generated_exports(
        root=export_root,
        now=now,
        preserve_storage_keys={normalized_key},
    )
    payload_path = export_root / normalized_key
    manifest_path = _manifest_path(payload_path)
    if (
        not manifest_path.is_file()
        or manifest_path.is_symlink()
        or manifest_path.resolve().parent != export_root
    ):
        raise GeneratedExportAccessError(404, "generated export not found")
    try:
        manifest = _load_manifest(manifest_path)
    except (OSError, ValueError, TypeError, json.JSONDecodeError) as exc:
        raise GeneratedExportAccessError(409, "generated export metadata failed validation") from exc

    presented_digest = hashlib.sha256(str(capability or "").encode("utf-8")).hexdigest()
    if not hmac.compare_digest(presented_digest, manifest["capabilityDigest"]):
        raise GeneratedExportAccessError(403, "generated export capability rejected")

    current_time = _as_utc(now or datetime.now(timezone.utc))
    if _parse_utc(manifest["expiresAt"]) <= current_time:
        _delete_export_pair(export_root, normalized_key)
        raise GeneratedExportAccessError(410, "generated export expired")

    open_flags = os.O_RDONLY | getattr(os, "O_CLOEXEC", 0)
    if hasattr(os, "O_NOFOLLOW"):
        open_flags |= os.O_NOFOLLOW
    try:
        descriptor = os.open(payload_path, open_flags)
    except OSError as exc:
        raise GeneratedExportAccessError(404, "generated export not found") from exc
    try:
        stream = os.fdopen(descriptor, "rb")
    except Exception:
        os.close(descriptor)
        raise
    try:
        file_status = os.fstat(stream.fileno())
        if not stat.S_ISREG(file_status.st_mode):
            raise GeneratedExportAccessError(404, "generated export not found")
        actual_sha256 = _sha256_stream(stream)
        if file_status.st_size != manifest["size"] or not hmac.compare_digest(actual_sha256, manifest["sha256"]):
            raise GeneratedExportAccessError(409, "generated export integrity check failed")
        stream.seek(0)
    except GeneratedExportAccessError:
        stream.close()
        raise
    except OSError as exc:
        stream.close()
        raise GeneratedExportAccessError(409, "generated export integrity check failed") from exc
    except Exception:
        stream.close()
        raise

    return GeneratedExportFile(
        stream=stream,
        storage_key=normalized_key,
        mime_type=manifest["mimeType"],
        sha256=manifest["sha256"],
        size=manifest["size"],
        created_at=manifest["createdAt"],
        expires_at=manifest["expiresAt"],
    )


def _manifest_path(payload_path: Path) -> Path:
    return payload_path.with_name(f"{payload_path.name}{_MANIFEST_SUFFIX}")


def _storage_key_from_manifest_name(name: str) -> str:
    if not str(name or "").endswith(_MANIFEST_SUFFIX):
        return ""
    storage_key = name[:-len(_MANIFEST_SUFFIX)]
    return storage_key if _is_valid_storage_key(storage_key) else ""


def _is_valid_storage_key(storage_key: str) -> bool:
    match = _STORAGE_KEY_PATTERN.fullmatch(str(storage_key or ""))
    if not match:
        return False
    try:
        return str(uuid.UUID(match.group("uuid"))) == match.group("uuid")
    except ValueError:
        return False


def _load_manifest(path: Path) -> Dict[str, Any]:
    value = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(value, dict) or set(value) != _MANIFEST_FIELDS:
        raise ValueError("invalid generated export manifest fields")
    capability_digest = str(value.get("capabilityDigest") or "")
    sha256 = str(value.get("sha256") or "")
    mime_type = str(value.get("mimeType") or "").strip()
    size = value.get("size")
    if not re.fullmatch(r"[0-9a-f]{64}", capability_digest):
        raise ValueError("invalid generated export capability digest")
    if not re.fullmatch(r"[0-9a-f]{64}", sha256):
        raise ValueError("invalid generated export digest")
    if isinstance(size, bool) or not isinstance(size, int) or size < 0:
        raise ValueError("invalid generated export size")
    if not re.fullmatch(r"[A-Za-z0-9!#$&^_.+-]+/[A-Za-z0-9!#$&^_.+-]+", mime_type):
        raise ValueError("invalid generated export MIME type")
    created_at = _format_utc(_parse_utc(value.get("createdAt")))
    expires_at = _format_utc(_parse_utc(value.get("expiresAt")))
    if _parse_utc(expires_at) <= _parse_utc(created_at):
        raise ValueError("invalid generated export expiry")
    return {
        "capabilityDigest": capability_digest,
        "sha256": sha256,
        "size": size,
        "mimeType": mime_type,
        "createdAt": created_at,
        "expiresAt": expires_at,
    }


def _parse_utc(value: Any) -> datetime:
    text = str(value or "").strip()
    if not text:
        raise ValueError("UTC timestamp is required")
    parsed = datetime.fromisoformat(text.replace("Z", "+00:00"))
    if parsed.tzinfo is None:
        raise ValueError("UTC timestamp must include timezone")
    return parsed.astimezone(timezone.utc)


def _format_utc(value: datetime) -> str:
    return _as_utc(value).isoformat(timespec="microseconds").replace("+00:00", "Z")


def _as_utc(value: datetime) -> datetime:
    if value.tzinfo is None:
        return value.replace(tzinfo=timezone.utc)
    return value.astimezone(timezone.utc)


def _sha256_file(path: Path) -> str:
    with path.open("rb") as stream:
        return _sha256_stream(stream)


def _sha256_stream(stream: BinaryIO) -> str:
    digest = hashlib.sha256()
    for chunk in iter(lambda: stream.read(1024 * 1024), b""):
        digest.update(chunk)
    return digest.hexdigest()


def _delete_export_pair(root: Path, storage_key: str) -> None:
    payload_path = Path(root) / storage_key
    _safe_unlink(payload_path)
    _safe_unlink(_manifest_path(payload_path))


def _safe_unlink(path: Optional[Path]) -> None:
    if path is None:
        return
    try:
        path.unlink(missing_ok=True)
    except OSError:
        pass


def _is_older_than_grace(path: Path, now: datetime, grace_seconds: float) -> bool:
    try:
        modified_at = datetime.fromtimestamp(path.stat().st_mtime, tz=timezone.utc)
    except OSError:
        return True
    return modified_at <= now - timedelta(seconds=grace_seconds)


def _parse_json_object(content: str) -> Dict[str, Any]:
    raw = str(content or "").strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```(?:json)?", "", raw, flags=re.IGNORECASE).strip()
        raw = re.sub(r"```$", "", raw).strip()
    try:
        parsed = json.loads(raw)
        return parsed if isinstance(parsed, dict) else {}
    except Exception:
        return {}


def _title_from_markdown(content: str) -> str:
    for line in str(content or "").splitlines():
        match = re.match(r"^#\s+(.+)$", line.strip())
        if match:
            return _clean_inline_markdown(match.group(1))[:60]
    return ""


def _title_from_metadata(metadata: Dict[str, Any], fallback: str) -> str:
    for key in ("sourceTitle", "title", "topic", "intent"):
        value = str(metadata.get(key) or "").strip()
        if value:
            return value[:60]
    agent = str(metadata.get("executedAgent") or metadata.get("targetAgent") or metadata.get("agentName") or "").strip()
    return _agent_title(agent) or fallback


def _agent_title(agent: str) -> str:
    titles = {
        "textbook_knowledge_agent": "教材知识整理",
        "meeting_summary_agent": "会议纪要",
        "meeting_resource_recommendation_agent": "学习资源推荐",
        "ppt_outline_agent": "PPT 大纲",
        "ppt_layout_agent": "PPT 布局方案",
        "ppt_review_agent": "PPT 审查报告",
    }
    if agent.startswith("textbook_question_"):
        return "题库导出"
    return titles.get(agent, "")


def _extract_options(question: Dict[str, Any]) -> List[Dict[str, Any]]:
    body = question.get("body") if isinstance(question.get("body"), dict) else {}
    options = body.get("options") or question.get("options") or []
    if not isinstance(options, list):
        return []
    normalized = []
    for option in options:
        if isinstance(option, dict):
            normalized.append({
                "key": str(option.get("key") or option.get("label") or "").strip(),
                "text": str(option.get("text") or option.get("content") or "").strip(),
            })
        else:
            normalized.append({"key": "", "text": str(option or "").strip()})
    return normalized


def _format_answer(answer: Any) -> str:
    if answer is None:
        return ""
    if isinstance(answer, dict):
        for key in (
            "correctOption",
            "correctOptions",
            "correctKey",
            "correctKeys",
            "correct",
            "blanks",
            "referenceAnswer",
            "answerPoints",
            "keyPoints",
            "finalAnswer",
            "steps",
            "proofSteps",
            "conclusion",
            "solutionOutline",
            "referenceSolution",
            "testCases",
            "expectedResult",
            "pairs",
            "orderedKeys",
            "value",
            "values",
        ):
            if key in answer:
                return _join_values(answer.get(key))
        return json.dumps(answer, ensure_ascii=False)
    return _join_values(answer)


def _join_values(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, list):
        return "、".join(str(item) for item in value if str(item).strip())
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False)
    return str(value).strip()


def _looks_like_markdown(content: str) -> bool:
    return bool(re.search(r"^#{1,6}\s+|^[-*]\s+|^\d+[.)]\s+", content, flags=re.MULTILINE))


def _extract_mermaid_code(content: str) -> str:
    match = re.search(r"```mermaid\s*([\s\S]*?)```", str(content or ""), flags=re.IGNORECASE)
    if match:
        return match.group(1).strip()
    stripped = str(content or "").strip()
    if re.match(r"^(mindmap|flowchart|graph|sequenceDiagram|classDiagram|stateDiagram|erDiagram|journey|gantt)\b", stripped, flags=re.IGNORECASE):
        return stripped
    return ""


def _clean_inline_markdown(text: str) -> str:
    value = str(text or "")
    value = re.sub(r"`([^`]+)`", r"\1", value)
    value = re.sub(r"\*\*([^*]+)\*\*", r"\1", value)
    value = re.sub(r"\*([^*]+)\*", r"\1", value)
    value = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", value)
    return value.strip()


def _safe_sheet_name(value: str) -> str:
    name = re.sub(r"[\[\]:*?/\\]", "", str(value or "Sheet1")).strip() or "Sheet1"
    return name[:31]


def _cell_text(value: Any) -> str:
    text = "" if value is None else str(value)
    return re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)[:32000]


def _column_name(index: int) -> str:
    result = ""
    value = index
    while value:
        value, remainder = divmod(value - 1, 26)
        result = chr(65 + remainder) + result
    return result or "A"


def _slugify(value: str) -> str:
    text = re.sub(r"[^\w\u4e00-\u9fff-]+", "-", str(value or "").strip(), flags=re.UNICODE)
    text = re.sub(r"-+", "-", text).strip("-")
    return text[:48] or "generated-content"


__all__ = [
    "EXPORT_MAX_BYTES",
    "EXPORT_ROOT",
    "EXPORT_STAGING_GRACE_SECONDS",
    "EXPORT_TTL_HOURS",
    "EXPORT_URL_PATH",
    "GeneratedExportAccessError",
    "GeneratedExportFile",
    "GeneratedExportResult",
    "cleanup_generated_exports",
    "export_generated_answer",
    "open_generated_export",
]
