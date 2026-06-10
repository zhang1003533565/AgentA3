from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List

from app.rag.indexing.multimodal_parser import MultimodalParser


@dataclass
class LoadedDocument:
    id: str
    content: str
    source: str


class DocumentLoader:
    SUPPORTED_SUFFIXES = {
        ".csv",
        ".docx",
        ".gif",
        ".htm",
        ".html",
        ".jpeg",
        ".jpg",
        ".json",
        ".markdown",
        ".md",
        ".pdf",
        ".png",
        ".tsv",
        ".txt",
        ".webp",
    }

    def __init__(self) -> None:
        self.multimodal_parser = MultimodalParser()

    def load(self, source: str) -> Iterable[LoadedDocument]:
        path = Path(source)
        if path.is_dir():
            return self.load_directory(path)
        if path.is_file() and path.suffix.lower() in self.SUPPORTED_SUFFIXES:
            return [self._load_file(path)]
        return []

    def load_directory(self, source: Path) -> List[LoadedDocument]:
        documents: List[LoadedDocument] = []
        if not source.exists():
            return documents
        for path in sorted(source.rglob("*")):
            if not path.is_file() or path.suffix.lower() not in self.SUPPORTED_SUFFIXES:
                continue
            documents.append(self._load_file(path))
        return documents

    def _load_file(self, path: Path) -> LoadedDocument:
        if path.suffix.lower() == ".docx":
            content = self._load_docx(path)
        elif path.suffix.lower() in {".txt"}:
            content = path.read_text(encoding="utf-8", errors="ignore")
        else:
            content = self.multimodal_parser.parse(str(path)).get("text", "")
        return LoadedDocument(
            id=str(path.resolve()),
            content=content,
            source=str(path),
        )

    def _load_docx(self, path: Path) -> str:
        from docx import Document

        document = Document(str(path))
        blocks: List[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                blocks.append(text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    blocks.append("\t".join(cells))
        return "\n".join(blocks)
