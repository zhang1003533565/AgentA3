# -*- coding: utf-8 -*-
"""Convert 作品报告 markdown to styled docx with figure placeholders."""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = Path(r"e:\zzs\github\AgentA3\docs\submission\作品报告-AgentA3.md")
OUT = Path(r"e:\zzs\github\AgentA3\docs\submission\作品报告-AgentA3.docx")
TPL = Path(r"c:\Users\zzs\Desktop\视频剪辑\文档\05-3 作品报告（人工智能挑战赛，2023版）模板.docx")


def set_run_font(run, name="宋体", size=12, bold=False, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    ascii_font = "Times New Roman" if name not in ("黑体", "Consolas") else name
    if name == "Consolas":
        ascii_font = "Consolas"
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), "宋体" if name == "Consolas" else name)
    if color:
        run.font.color.rgb = color


def add_heading_cn(doc, text, level):
    p = doc.add_paragraph()
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, "黑体", 16, True)
    elif level == 1:
        run = p.add_run(text)
        set_run_font(run, "黑体", 14, True)
    else:
        run = p.add_run(text)
        set_run_font(run, "黑体", 12, True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_body(doc, text, first_indent=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, "宋体", 12)
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(3)
    if first_indent:
        pf.first_line_indent = Cm(0.74)
    return p


def add_quote_box(doc, lines):
    p = doc.add_paragraph()
    run = p.add_run("\n".join(lines))
    set_run_font(run, "楷体", 10.5, color=RGBColor(0x55, 0x55, 0x55))
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.line_spacing = 1.3
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F0F0F0")
    shd.set(qn("w:val"), "clear")
    p._element.get_or_add_pPr().append(shd)
    return p


def add_figure_placeholder(doc, title, hint_lines):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_run_font(run, "宋体", 10.5, True, RGBColor(0x33, 0x33, 0x33))

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Cm(14)
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "D9D9D9")
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "dashed")
        b.set(qn("w:sz"), "12")
        b.set(qn("w:color"), "808080")
        tc_borders.append(b)
    tc_pr.append(tc_borders)

    cell.text = ""
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run("【截图占位】\n" + "\n".join(hint_lines))
    set_run_font(r, "宋体", 10.5, color=RGBColor(0x66, 0x66, 0x66))
    for _ in range(4):
        cell.add_paragraph()
    doc.add_paragraph()


def add_code_block(doc, code):
    for line in code.splitlines() or [""]:
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        set_run_font(run, "Consolas", 9)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F5F5F5")
        shd.set(qn("w:val"), "clear")
        p._element.get_or_add_pPr().append(shd)
    doc.add_paragraph()


def add_table_from_rows(doc, rows):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            val = row[j] if j < len(row) else ""
            run = p.add_run(val)
            set_run_font(run, "宋体", 10.5, bold=(i == 0))
            if i == 0:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "E7E6E6")
                shd.set(qn("w:val"), "clear")
                cell._tc.get_or_add_tcPr().append(shd)
    doc.add_paragraph()


def main():
    text = SRC.read_text(encoding="utf-8")
    try:
        doc = Document(str(TPL))
        body = doc.element.body
        for child in list(body):
            if child.tag != qn("w:sectPr"):
                body.remove(child)
    except Exception as exc:
        print("template load failed", exc)
        doc = Document()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("中国计算机学会青少年\n人工智能挑战赛·作品报告")
    set_run_font(r, "黑体", 18, True)

    doc.add_paragraph()
    for k, v in [
        ("作品编号：", "（填写）"),
        ("作品名称：", "AgentA3——智慧校园个性化学习多智能体系统"),
        ("编写日期：", "2026-09-04"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(k)
        set_run_font(r1, "宋体", 14, True)
        r2 = p.add_run(v)
        set_run_font(r2, "宋体", 14)

    doc.add_page_break()

    lines = text.splitlines()
    start = 0
    for idx, line in enumerate(lines):
        if line.startswith("## 第1章") or line.startswith("# 第1章"):
            start = idx
            break

    i = start
    in_code = False
    code_buf = []
    table_buf = []
    in_table = False
    quote_buf = []
    in_quote = False
    fig_buf = []
    in_fig = False

    def flush_table():
        nonlocal table_buf, in_table
        if not table_buf:
            in_table = False
            return
        rows = []
        for row in table_buf:
            if re.match(r"^\|?\s*-+", row):
                continue
            cells = [c.strip() for c in row.strip().strip("|").split("|")]
            rows.append(cells)
        add_table_from_rows(doc, rows)
        table_buf = []
        in_table = False

    def flush_quote():
        nonlocal quote_buf, in_quote
        if quote_buf:
            add_quote_box(doc, quote_buf)
        quote_buf = []
        in_quote = False

    def flush_fig():
        nonlocal fig_buf, in_fig
        if not fig_buf:
            in_fig = False
            return
        title = fig_buf[0]
        hints = []
        for x in fig_buf[1:]:
            x = x.lstrip("> ").strip()
            if x:
                hints.append(x)
        add_figure_placeholder(doc, title, hints or ["请在此插入截图"])
        fig_buf = []
        in_fig = False

    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            if in_code:
                add_code_block(doc, "\n".join(code_buf))
                code_buf = []
                in_code = False
            else:
                flush_table()
                flush_quote()
                flush_fig()
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if line.startswith("**【图") or (line.startswith("【图") and "截图位】" in line):
            flush_table()
            flush_quote()
            flush_fig()
            in_fig = True
            fig_buf = [re.sub(r"\*\*", "", line).strip()]
            i += 1
            continue
        if in_fig:
            if line.startswith(">") or line.strip() == "":
                if line.startswith(">"):
                    fig_buf.append(line)
                elif line.strip() == "" and i + 1 < len(lines) and not lines[i + 1].startswith(">"):
                    flush_fig()
                i += 1
                continue
            flush_fig()
            continue

        if line.startswith(">"):
            if not in_quote:
                flush_table()
                flush_fig()
                in_quote = True
                quote_buf = []
            quote_buf.append(line.lstrip("> ").strip())
            i += 1
            continue
        if in_quote:
            flush_quote()

        if line.strip().startswith("|"):
            if not in_table:
                flush_fig()
                flush_quote()
                in_table = True
                table_buf = []
            table_buf.append(line)
            i += 1
            continue
        if in_table:
            flush_table()

        if line.startswith("#### "):
            add_heading_cn(doc, line[5:].strip(), 3)
        elif line.startswith("### "):
            add_heading_cn(doc, line[4:].strip(), 2)
        elif line.startswith("## "):
            t = line[3:].strip()
            if t.startswith("第") and "章" in t and t != "第1章 作品概述":
                doc.add_page_break()
            add_heading_cn(doc, t, 1)
        elif line.startswith("# "):
            add_heading_cn(doc, line[2:].strip(), 0)
        elif line.strip() in ("---", ""):
            pass
        else:
            content = line.strip()
            content = re.sub(r"\*\*(.+?)\*\*", r"\1", content)
            content = re.sub(r"`([^`]+)`", r"\1", content)
            if content.startswith("- "):
                p = doc.add_paragraph()
                run = p.add_run("• " + content[2:])
                set_run_font(run, "宋体", 12)
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.left_indent = Cm(0.74)
            elif re.match(r"^\d+\.\s", content):
                p = doc.add_paragraph()
                run = p.add_run(content)
                set_run_font(run, "宋体", 12)
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.first_line_indent = Cm(0.74)
            else:
                add_body(doc, content)
        i += 1

    flush_table()
    flush_quote()
    flush_fig()
    if in_code and code_buf:
        add_code_block(doc, "\n".join(code_buf))

    doc.save(str(OUT))
    print("saved", OUT, "size", OUT.stat().st_size)


if __name__ == "__main__":
    main()
