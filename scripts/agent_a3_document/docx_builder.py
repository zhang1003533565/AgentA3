"""Build the editable AgentA3 project report from reviewed Markdown blocks."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path, PurePosixPath
import re
from tempfile import TemporaryDirectory
from typing import Sequence

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

from .source_loader import Block, extract_evidence_rows, load_source
from .styles import (
    BLUE_GRAY,
    CJK_SERIF_FONT,
    CJK_SANS_FONT,
    DARK,
    LATIN_FONT,
    LIGHT_GRAY,
    MONO_FONT,
    PRIMARY,
    RISK,
    USABLE_PAGE_WIDTH_DXA,
    WARNING,
    add_numbering_instance,
    add_toc_field,
    apply_list_numbering,
    configure_cover_section,
    configure_document,
    configure_running_section,
    configure_table,
    shade_paragraph,
)


_FIGURE_ALIASES = {
    "tech-profile-evidence-state.png": "profile-evidence-flow.png",
    "tech-leader-routing.png": "agent-capability-groups.png",
    "tech-xfyun-asr-websocket.png": "meeting-asr-loop.png",
    "tech-meeting-postprocess-sequence.png": "meeting-asr-loop.png",
    "tech-maxkb-grounded-answer.png": "maxkb-grounding-flow.png",
    "tech-resource-envelope.png": "resource-envelope.png",
    "tech-question-paper-trust-chain.png": "question-paper-exam-loop.png",
    "tech-exam-consistency.png": "question-paper-exam-loop.png",
}

_DESIGN_IMAGES = {
    "profile-radar-design.png": Path(
        "docs/designs/profile-radar/profile-radar-design.png"
    ),
    "ai-conversation-design.png": Path(
        "docs/designs/ai-conversation/ai-conversation-design.png"
    ),
    "meeting-home-design.png": Path(
        "docs/designs/meeting-home/meeting-home-design.png"
    ),
}

_SUPPLEMENTAL_FIGURES = {
    "3.2 系统上下文": (
        "system-context.png",
        "系统上下文与受控外部依赖",
    ),
    "3.3 四部分总体架构": (
        "four-part-architecture.png",
        "移动端、Web、Java 与 Python 四部分总体架构",
    ),
    "8.3 需求—设计—接口—测试追踪矩阵": (
        "requirements-trace.png",
        "需求、设计、接口与测试追踪关系",
    ),
}

_EXPECTED_GENERATED_IMAGES = {
    "agent-capability-groups.png",
    "core-entity-relations.png",
    "deployment-boundary.png",
    "four-part-architecture.png",
    "maxkb-grounding-flow.png",
    "meeting-asr-loop.png",
    "profile-evidence-flow.png",
    "question-paper-exam-loop.png",
    "requirements-trace.png",
    "resource-envelope.png",
    "system-context.png",
}

_INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
_EXISTING_CAPTION_RE = re.compile(r"^图(?:\d+|[一二三四五六七八九十]+)-\d+\s*")
_FUNCTION_HEADING_RE = re.compile(r"^\d+\.\d+\s+(FUNC-\d+)\s+")
_FUNCTION_DETAIL_RE = re.compile(
    r"^\*\*(?P<label>[^*]+)\*\*[：:]\s*(?P<value>.*)$",
    re.DOTALL,
)
_FUNCTION_DETAIL_LABELS = (
    "目标",
    "参与者",
    "触发条件",
    "前置条件",
    "主流程",
    "替代流程",
    "异常流程",
    "后置条件",
    "业务规则",
    "涉及接口",
    "数据实体",
    "验收条件",
    "测试编号",
)
_WINDOWS_DRIVE_RE = re.compile(r"^[A-Za-z]:")
_WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


@dataclass(frozen=True)
class TocEntry:
    """One visible cached TOC row retained inside the native TOC field."""

    level: int
    title: str
    page: int | str


def materialize_toc_cache(
    docx_path: Path,
    entries: Sequence[TocEntry],
) -> Path:
    """Replace a TOC field's placeholder result with visible dotted entries."""

    path = Path(docx_path)
    normalized: list[TocEntry] = []
    for entry in entries:
        title = " ".join(entry.title.replace("\t", " ").split())
        if not title:
            raise ValueError("TOC entry title must not be empty")
        if entry.level < 0:
            raise ValueError("TOC entry level must be non-negative")
        if isinstance(entry.page, bool):
            raise ValueError("TOC entry page label must be a positive page number")
        if isinstance(entry.page, int):
            if entry.page <= 0:
                raise ValueError("TOC entry page number must be positive")
            page_label = str(entry.page)
        else:
            page_label = " ".join(entry.page.replace("\t", " ").split())
            if not page_label:
                raise ValueError("TOC entry page label must not be empty")
        normalized.append(TocEntry(min(entry.level, 2), title, page_label))
    if not normalized:
        raise ValueError("at least one TOC entry is required")

    document = Document(path)
    toc_paragraph = next(
        (
            paragraph
            for paragraph in document.paragraphs
            if any(
                (node.text or "").strip().startswith("TOC ")
                for node in paragraph._p.findall(
                    ".//w:instrText", {"w": _WORD_NS}
                )
            )
        ),
        None,
    )
    if toc_paragraph is None:
        raise ValueError("DOCX does not contain a native TOC field")

    try:
        cache_style = document.styles["TOC Cache"]
    except KeyError:
        cache_style = document.styles.add_style("TOC Cache", WD_STYLE_TYPE.PARAGRAPH)
        cache_style.font.name = LATIN_FONT
        cache_style.font.size = Pt(8.5)

    next_element = toc_paragraph._p.getnext()
    while next_element is not None:
        style = next_element.find(qn("w:pPr"))
        style = style.find(qn("w:pStyle")) if style is not None else None
        if style is None or style.get(qn("w:val")) != cache_style.style_id:
            break
        following = next_element.getnext()
        next_element.getparent().remove(next_element)
        next_element = following

    toc_paragraph.clear()
    toc_paragraph.style = cache_style
    field_run = toc_paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    field_run._r.extend((begin, instruction, separate))

    paragraphs = [toc_paragraph]
    for _ in normalized[1:]:
        element = OxmlElement("w:p")
        paragraphs[-1]._p.addnext(element)
        paragraph = Paragraph(element, toc_paragraph._parent)
        paragraph.style = cache_style
        paragraphs.append(paragraph)

    for paragraph, entry in zip(paragraphs, normalized):
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = Cm(0.55 * entry.level)
        paragraph_format.right_indent = None
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.line_spacing = 1.0
        paragraph_format.keep_together = True
        paragraph_format.tab_stops.add_tab_stop(
            Cm(16.1), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )

        title_run = paragraph.add_run(entry.title)
        title_run.font.size = Pt(8.5 if entry.level == 0 else 8.0)
        title_run.bold = bool(re.match(r"^第[一二三四五六七八九十]+章", entry.title))
        _set_run_fonts(title_run, east_asia=CJK_SANS_FONT)
        tab_run = paragraph.add_run("\t")
        _set_run_fonts(tab_run, east_asia=CJK_SANS_FONT)
        page_run = paragraph.add_run(str(entry.page))
        page_run.font.size = title_run.font.size
        _set_run_fonts(page_run, east_asia=CJK_SANS_FONT)

    end_run = paragraphs[-1].add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)

    with TemporaryDirectory(prefix=".agent-a3-toc-", dir=path.parent) as temporary_dir:
        temporary_path = Path(temporary_dir) / path.name
        document.save(temporary_path)
        temporary_path.replace(path)
    return path


def build_document(
    source_dir: Path,
    evidence_path: Path,
    asset_dir: Path,
    output_path: Path,
) -> Path:
    """Build a self-contained DOCX and return ``output_path``.

    All source and asset validation is completed before an output directory is
    created, so missing or unsafe assets cannot leave a partial deliverable.
    """

    builder = _DocumentBuilder(
        source_dir=Path(source_dir),
        evidence_path=Path(evidence_path),
        asset_dir=Path(asset_dir),
        output_path=Path(output_path),
    )
    return builder.build()


class _DocumentBuilder:
    def __init__(
        self,
        *,
        source_dir: Path,
        evidence_path: Path,
        asset_dir: Path,
        output_path: Path,
    ) -> None:
        self.source_dir = source_dir
        self.evidence_path = evidence_path
        self.asset_dir = asset_dir
        self.output_path = output_path
        self.blocks: list[Block] = []
        self.figure_assets: dict[str, Path] = {}
        self.supplemental_assets: dict[str, Path] = {}
        self.repository_root = Path()

        self.document = Document()
        self.numbering = None
        self.front_matter_started = False
        self.cover_mode = True
        self.chapter_index = 0
        self.current_chapter = 0
        self.current_section_title = ""
        self.figure_counter = 0
        self.table_counter = 0
        self.pending_supplement: tuple[str, str] | None = None
        self.previous_block_kind = ""
        self.current_decimal_num_id: int | None = None
        self.cover_paragraph_index = 0

    def build(self) -> Path:
        self._preflight()
        self._initialize_document()
        index = 0
        while index < len(self.blocks):
            block = self.blocks[index]
            self._render_block(block)
            self.previous_block_kind = block.kind
            function_match = (
                _FUNCTION_HEADING_RE.match(block.text)
                if block.kind == "heading" and block.level == 2
                else None
            )
            if function_match is not None:
                detail_start = index + 1
                detail_end = detail_start
                while detail_end < len(self.blocks):
                    candidate = self.blocks[detail_end]
                    if candidate.kind == "section_break" or (
                        candidate.kind == "heading" and candidate.level <= block.level
                    ):
                        break
                    detail_end += 1
                self._render_function_details(
                    function_match.group(1),
                    self.blocks[detail_start:detail_end],
                )
                self.previous_block_kind = "table"
                index = detail_end
                continue
            index += 1
        if self.pending_supplement is not None:
            self._render_pending_supplement()

        self.output_path.parent.mkdir(parents=True, exist_ok=True)
        with TemporaryDirectory(
            prefix=".agent-a3-docx-", dir=self.output_path.parent
        ) as temporary_dir:
            temporary_path = Path(temporary_dir) / self.output_path.name
            self.document.save(temporary_path)
            temporary_path.replace(self.output_path)
        return self.output_path

    def _preflight(self) -> None:
        self.blocks = load_source(self.source_dir)
        evidence_rows = extract_evidence_rows(self.evidence_path)
        if not evidence_rows or any(
            not value.strip()
            for row in evidence_rows
            for value in (
                row.id,
                row.claim,
                row.status,
                row.evidence,
                row.final_wording,
            )
        ):
            raise ValueError("evidence index contains an empty required field")

        self.repository_root = _find_repository_root(
            self.source_dir, self.evidence_path, self.asset_dir
        ).resolve()
        asset_root = self.asset_dir.resolve()
        generated: dict[str, Path] = {}
        for path in self.asset_dir.glob("*.png"):
            if not path.is_file():
                continue
            generated[path.name] = _resolve_within(
                path,
                asset_root,
                error_message=f"generated asset escapes asset directory: {path.name}",
            )
        missing_generated = sorted(_EXPECTED_GENERATED_IMAGES - generated.keys())
        if missing_generated:
            raise ValueError(
                "missing figure asset: " + ", ".join(missing_generated)
            )

        for block in self.blocks:
            if block.kind != "figure":
                continue
            source = block.attrs["src"]
            self.figure_assets[source] = self._resolve_figure(source, generated)

        for heading, (name, _caption) in _SUPPLEMENTAL_FIGURES.items():
            path = generated.get(name)
            if path is None or not path.is_file():
                raise ValueError(f"missing figure asset: {name}")
            self.supplemental_assets[heading] = path

    def _resolve_figure(self, source: str, generated: dict[str, Path]) -> Path:
        normalized = PurePosixPath(source)
        if (
            normalized.is_absolute()
            or ".." in normalized.parts
            or source.startswith("~")
            or _WINDOWS_DRIVE_RE.match(source)
            or "://" in source
            or "\\" in source
        ):
            raise ValueError(
                f"repository-relative figure path required: {source}"
            )

        name = normalized.name
        generated_name = _FIGURE_ALIASES.get(name, name)
        candidate = generated.get(generated_name)
        if candidate is None:
            repository_relative = _DESIGN_IMAGES.get(name, Path(*normalized.parts))
            repository_candidate = self.repository_root / repository_relative
            if repository_candidate.is_file():
                candidate = _resolve_within(
                    repository_candidate,
                    self.repository_root,
                    error_message=(
                        f"repository-relative figure path escapes repository: {source}"
                    ),
                )
        if candidate is None or not candidate.is_file():
            raise ValueError(f"missing figure asset: {source}")
        return candidate

    def _initialize_document(self) -> None:
        self.numbering = configure_document(self.document)
        configure_cover_section(self.document.sections[0])
        properties = self.document.core_properties
        properties.title = (
            "AgentA3 智慧校园个性化学习多智能体系统——需求、设计与开发说明书"
        )
        properties.subject = "需求、设计、开发、测试与部署说明"
        properties.author = "AgentA3 项目组"
        properties.keywords = "AgentA3, 个性化学习, 多智能体, 工程证据"
        properties.comments = "由受审阅 Markdown 源和 48 条工程证据生成"

    def _render_block(self, block: Block) -> None:
        if block.kind == "section_break":
            self._handle_section_break(block.text)
        elif block.kind == "heading":
            self._render_heading(block)
        elif block.kind == "paragraph":
            self._render_paragraph(block.text)
            self._render_pending_supplement()
        elif block.kind in {"bullet", "numbered"}:
            self._render_list_item(block)
        elif block.kind == "table":
            self._render_table(block)
        elif block.kind == "figure":
            self._render_figure(
                self.figure_assets[block.attrs["src"]],
                block.attrs["caption"],
                float(block.attrs["width_cm"]),
            )
        elif block.kind == "callout":
            self._render_callout(block)
        elif block.kind == "code":
            self._render_code(block)
        elif block.kind == "toc":
            self._render_toc()
        elif block.kind == "page_break":
            self.document.add_page_break()
        else:
            raise ValueError(f"unsupported block kind: {block.kind}")

    def _handle_section_break(self, title: str) -> None:
        self.current_section_title = title
        self.figure_counter = 0
        self.table_counter = 0
        self.current_decimal_num_id = None
        if title == "前置材料":
            self.current_chapter = 0
            return

        if not self.front_matter_started:
            self._start_front_matter()
        self.chapter_index += 1
        self.current_chapter = self.chapter_index
        section = self.document.add_section(WD_SECTION.NEW_PAGE)
        configure_running_section(
            section,
            title,
            page_format="decimal",
            page_start=1 if self.chapter_index == 1 else None,
        )
        self.cover_mode = False

    def _start_front_matter(self) -> None:
        if self.front_matter_started:
            return
        section = self.document.add_section(WD_SECTION.NEW_PAGE)
        configure_running_section(
            section,
            "前置材料",
            page_format="upperRoman",
            page_start=1,
        )
        self.front_matter_started = True
        self.cover_mode = False
        self.current_chapter = 0
        self.figure_counter = 0
        self.table_counter = 0

    def _render_heading(self, block: Block) -> None:
        if self.cover_mode and block.level >= 2:
            self._start_front_matter()

        if self.cover_mode:
            style = "Title" if block.level == 1 else "Heading 1"
        elif block.level == 1:
            style = "Chapter Title"
        else:
            style = f"Heading {min(block.level - 1, 3)}"
        paragraph = self.document.add_paragraph(style=style)
        if self.cover_mode and block.level == 1 and "——" in block.text:
            system_title, separator, document_scope = block.text.partition("——")
            self._add_inline_text(paragraph, system_title)
            paragraph.add_run().add_break()
            self._add_inline_text(paragraph, f"{separator}{document_scope}")
        else:
            self._add_inline_text(paragraph, block.text)
        if self.cover_mode:
            paragraph.paragraph_format.space_before = Pt(42)
            paragraph.paragraph_format.space_after = Pt(24)

        supplemental = _SUPPLEMENTAL_FIGURES.get(block.text)
        if supplemental:
            self.pending_supplement = supplemental

    def _render_paragraph(self, text: str) -> None:
        style = "Subtitle" if self.cover_mode else "Normal"
        paragraph = self.document.add_paragraph(style=style)
        self._add_inline_text(paragraph, text)
        if self.cover_mode:
            self.cover_paragraph_index += 1
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.space_after = Pt(
                16 if self.cover_paragraph_index < 3 else 7
            )

    def _render_list_item(self, block: Block) -> None:
        paragraph = self.document.add_paragraph(style="Normal")
        paragraph.paragraph_format.first_line_indent = None
        self._add_inline_text(paragraph, block.text)
        level = max(0, block.level - 1)
        if block.kind == "bullet":
            num_id = self.numbering.bullet_num_id
        else:
            if self.previous_block_kind != "numbered":
                self.current_decimal_num_id = add_numbering_instance(
                    self.document, self.numbering.decimal_abstract_id
                )
            num_id = self.current_decimal_num_id
        apply_list_numbering(paragraph, num_id=num_id, level=level)

    def _render_table(self, block: Block) -> None:
        self.table_counter += 1
        header = block.rows[0]
        label = "与".join(cell for cell in (header[0], header[-1]) if cell)
        self._add_caption("表", label, above=True)

        table = self.document.add_table(
            rows=len(block.rows), cols=len(block.rows[0])
        )
        if block.attrs.get("layout") == "function-detail":
            label_width = round(USABLE_PAGE_WIDTH_DXA * 0.18)
            widths = (label_width, USABLE_PAGE_WIDTH_DXA - label_width)
        else:
            widths = _table_column_widths(
                block.rows, total_dxa=USABLE_PAGE_WIDTH_DXA
            )
        configure_table(table, widths)
        for row_index, (row, source_row) in enumerate(zip(table.rows, block.rows)):
            for column_index, (cell, text) in enumerate(zip(row.cells, source_row)):
                paragraph = cell.paragraphs[0]
                paragraph.clear()
                paragraph.style = "Table Text"
                self._add_inline_text(paragraph, text)
                if row_index == 0 or (
                    block.attrs.get("layout") == "function-detail"
                    and column_index == 0
                ):
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        _set_run_fonts(run, east_asia=CJK_SANS_FONT)

    def _render_function_details(
        self,
        function_id: str,
        blocks: Sequence[Block],
    ) -> None:
        """Render repeated atomic-function fields as one compact traceability table."""

        rows: list[tuple[str, str]] = []
        current_label = ""
        current_parts: list[str] = []
        numbered_index = 0

        def flush() -> None:
            nonlocal current_label, current_parts, numbered_index
            if not current_label:
                return
            rows.append((current_label, "\n".join(current_parts).strip()))
            current_label = ""
            current_parts = []
            numbered_index = 0

        for block in blocks:
            if block.kind == "paragraph":
                match = _FUNCTION_DETAIL_RE.match(block.text)
                if match is None:
                    raise ValueError(
                        f"{function_id} contains an unlabeled detail paragraph"
                    )
                flush()
                current_label = match.group("label").strip()
                value = match.group("value").strip()
                if value:
                    current_parts.append(value)
            elif block.kind == "numbered":
                if not current_label:
                    raise ValueError(f"{function_id} contains an orphan numbered item")
                numbered_index += 1
                current_parts.append(f"{numbered_index}. {block.text}")
            elif block.kind == "bullet":
                if not current_label:
                    raise ValueError(f"{function_id} contains an orphan bullet item")
                current_parts.append(f"• {block.text}")
            else:
                raise ValueError(
                    f"{function_id} contains unsupported detail block: {block.kind}"
                )
        flush()

        labels = tuple(label for label, _ in rows)
        if labels != _FUNCTION_DETAIL_LABELS or any(not value for _, value in rows):
            raise ValueError(f"{function_id} has an incomplete detail contract")
        table_block = Block(
            kind="table",
            attrs={"layout": "function-detail"},
            rows=((function_id, "原子功能设计"), *rows),
        )
        self._render_table(table_block)

    def _render_figure(self, path: Path, caption: str, width_cm: float) -> None:
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(2)
        with Image.open(path) as image:
            pixel_width, pixel_height = image.size
        target_width = min(width_cm, 16.2)
        target_height = target_width * pixel_height / pixel_width
        if target_height > 20.5:
            target_height = 20.5
            target_width = target_height * pixel_width / pixel_height
        paragraph.add_run().add_picture(
            str(path), width=Cm(target_width), height=Cm(target_height)
        )
        cleaned_caption = _EXISTING_CAPTION_RE.sub("", caption).strip()
        self._add_caption("图", cleaned_caption, above=False)

    def _render_callout(self, block: Block) -> None:
        paragraph = self.document.add_paragraph(style="Callout")
        callout_type = block.attrs.get("type", "note")
        border = {
            "risk": RISK,
            "evidence": PRIMARY,
            "note": WARNING,
        }.get(callout_type, PRIMARY)
        title = paragraph.add_run(f"{block.attrs.get('title', '说明')}：")
        title.bold = True
        title.font.color.rgb = RGBColor.from_string(border)
        _set_run_fonts(title, east_asia=CJK_SANS_FONT)
        self._add_inline_text(paragraph, block.text)
        shade_paragraph(paragraph, LIGHT_GRAY, border_color=border)

    def _render_code(self, block: Block) -> None:
        paragraph = self.document.add_paragraph(style="Code")
        paragraph.paragraph_format.first_line_indent = None
        if block.attrs.get("language"):
            language = paragraph.add_run(f"[{block.attrs['language']}]\n")
            language.bold = True
            language.font.color.rgb = RGBColor.from_string(PRIMARY)
            _set_run_fonts(language, east_asia=CJK_SANS_FONT)
        run = paragraph.add_run(block.text)
        run.font.name = MONO_FONT
        run.font.size = Pt(8.5)
        _set_run_fonts(run, latin=MONO_FONT, east_asia=CJK_SANS_FONT)
        shade_paragraph(paragraph, LIGHT_GRAY, border_color=BLUE_GRAY)

    def _render_toc(self) -> None:
        if not self.front_matter_started:
            self._start_front_matter()
        title = self.document.add_paragraph("目录", style="Heading 1")
        title.paragraph_format.page_break_before = True
        toc = self.document.add_paragraph(style="Normal")
        add_toc_field(toc)

    def _render_pending_supplement(self) -> None:
        if self.pending_supplement is None:
            return
        name, caption = self.pending_supplement
        path = self.supplemental_assets.get(self._supplement_heading(name))
        if path is None:
            path = (self.asset_dir / name).resolve()
        self._render_figure(path, caption, 15.5)
        self.pending_supplement = None

    @staticmethod
    def _supplement_heading(name: str) -> str:
        return next(
            heading
            for heading, (filename, _caption) in _SUPPLEMENTAL_FIGURES.items()
            if filename == name
        )

    def _add_caption(self, prefix: str, label: str, *, above: bool) -> None:
        if prefix == "图":
            self.figure_counter += 1
            counter = self.figure_counter
        else:
            counter = self.table_counter
        paragraph = self.document.add_paragraph(style="Caption")
        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.LEFT if above else WD_ALIGN_PARAGRAPH.CENTER
        )
        paragraph.paragraph_format.keep_with_next = above
        chapter = self.current_chapter
        self._add_inline_text(
            paragraph,
            f"{prefix}{chapter}-{counter} {label}".strip(),
        )

    def _add_inline_text(self, paragraph, text: str) -> None:
        parts = _INLINE_RE.split(text)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
                _set_run_fonts(run, east_asia=CJK_SANS_FONT)
            elif part.startswith("`") and part.endswith("`"):
                run = paragraph.add_run(part[1:-1])
                run.font.name = MONO_FONT
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor.from_string(DARK)
                _set_run_fonts(run, latin=MONO_FONT, east_asia=CJK_SANS_FONT)
                _shade_run(run, LIGHT_GRAY)
            else:
                lines = part.split("\n")
                for index, line in enumerate(lines):
                    if index:
                        paragraph.add_run().add_break()
                    run = paragraph.add_run(line)
                    _set_run_fonts(run)


def _find_repository_root(*paths: Path) -> Path:
    for path in paths:
        resolved = path.resolve()
        candidates = (resolved, *resolved.parents) if resolved.is_dir() else resolved.parents
        for candidate in candidates:
            if (candidate / ".git").exists():
                return candidate
    raise ValueError("repository root could not be resolved for figure assets")


def _table_column_widths(
    rows: tuple[tuple[str, ...], ...], *, total_dxa: int
) -> tuple[int, ...]:
    weights = []
    for column in range(len(rows[0])):
        longest = max(len(row[column]) for row in rows)
        weights.append(max(6, min(longest, 28)))
    total_weight = sum(weights)
    widths = [total_dxa * weight // total_weight for weight in weights]
    remainders = [total_dxa * weight % total_weight for weight in weights]
    missing = total_dxa - sum(widths)
    order = sorted(
        range(len(widths)), key=lambda index: (-remainders[index], index)
    )
    for index in order[:missing]:
        widths[index] += 1
    return tuple(widths)


def _resolve_within(path: Path, boundary: Path, *, error_message: str) -> Path:
    resolved = path.resolve(strict=True)
    try:
        resolved.relative_to(boundary.resolve(strict=True))
    except ValueError as exc:
        raise ValueError(error_message) from exc
    return resolved


def _set_run_fonts(
    run,
    *,
    latin: str = LATIN_FONT,
    east_asia: str = CJK_SERIF_FONT,
) -> None:
    run.font.name = latin
    r_fonts = run._r.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:cs"), latin)


def _shade_run(run, fill: str) -> None:
    r_pr = run._r.get_or_add_rPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    r_pr.append(shading)
