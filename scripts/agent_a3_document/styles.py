"""Word styling and low-level OOXML helpers for the AgentA3 report."""

from __future__ import annotations

from dataclasses import dataclass

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


LATIN_FONT = "Times New Roman"
CJK_SERIF_FONT = "Source Han Serif SC"
CJK_SANS_FONT = "Source Han Sans SC"
MONO_FONT = "Courier New"

PRIMARY = "0F766E"
DARK = "123B45"
BLUE_GRAY = "4F6B7A"
LIGHT_HEADER = "E8F3F1"
LIGHT_GRAY = "F2F4F5"
WARNING = "B7791F"
RISK = "B42318"
USABLE_PAGE_WIDTH_DXA = round(16.6 * 1440 / 2.54)


@dataclass(frozen=True)
class NumberingDefinition:
    bullet_num_id: int
    decimal_abstract_id: int


def configure_document(document: Document) -> NumberingDefinition:
    """Apply the approved style system and return native list definitions."""

    _configure_styles(document)
    _set_document_defaults(document)
    set_update_fields(document)
    return _create_list_numbering(document)


def configure_section_geometry(section) -> None:
    """Apply the invariant A4 portrait geometry to a section."""

    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.0)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)


def configure_cover_section(section) -> None:
    """Keep the cover free of running furniture and page fields."""

    configure_section_geometry(section)
    sect_pr = section._sectPr
    for tag in ("w:headerReference", "w:footerReference", "w:pgNumType"):
        for element in list(sect_pr.findall(qn(tag))):
            sect_pr.remove(element)


def configure_running_section(
    section,
    chapter_title: str,
    *,
    page_format: str = "decimal",
    page_start: int | None = None,
) -> None:
    """Create an independent running header, page field, and numbering format."""

    configure_section_geometry(section)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    _clear_story(section.header)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(3)
    header.paragraph_format.tab_stops.add_tab_stop(Cm(16.6))
    left = header.add_run("AgentA3 · 智慧校园个性化学习")
    _format_run(left, size=8.5, color=BLUE_GRAY, east_asia=CJK_SANS_FONT)
    header.add_run("\t")
    right = header.add_run(chapter_title)
    _format_run(right, size=8.5, color=DARK, east_asia=CJK_SANS_FONT)
    _set_paragraph_bottom_border(header, PRIMARY, size="6")

    _clear_story(section.footer)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(2)
    add_field(footer, "PAGE", placeholder="1")
    for run in footer.runs:
        _format_run(run, size=9, color=BLUE_GRAY, east_asia=CJK_SERIF_FONT)

    _set_page_numbering(section, page_format, page_start)


def add_toc_field(paragraph) -> None:
    """Insert an updateable Word table-of-contents field."""

    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = None
    add_field(
        paragraph,
        'TOC \\o "1-3" \\h \\z \\u',
        placeholder="目录将在 Word 中更新",
    )


def add_field(paragraph, instruction: str, *, placeholder: str = "") -> None:
    """Append a complex Word field to a paragraph."""

    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction_text, separate, text, end))


def set_update_fields(document: Document) -> None:
    """Tell Word to refresh fields when the document opens."""

    settings = document.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def add_numbering_instance(document: Document, abstract_num_id: int) -> int:
    """Create a fresh numbering instance so a numbered list restarts at one."""

    numbering = document.part.numbering_part.element
    num_id = _next_numeric_attribute(numbering, "w:num", "w:numId")
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_num_id))
    num.append(abstract)
    numbering.append(num)
    return num_id


def apply_list_numbering(paragraph, *, num_id: int, level: int) -> None:
    """Attach a paragraph to a native Word numbering definition."""

    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(max(0, min(level, 8))))
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num))


def configure_table(table, column_widths_dxa: tuple[int, ...]) -> None:
    """Apply fixed layout, borders, repeating header, and cell margins."""

    if not column_widths_dxa or any(width <= 0 for width in column_widths_dxa):
        raise ValueError("table column widths must be positive")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    table_width = tbl_pr.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        _insert_before_first(
            tbl_pr,
            table_width,
            ("w:tblInd", "w:tblBorders", "w:tblLayout", "w:tblLook"),
        )
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), str(sum(column_widths_dxa)))

    table_indent = tbl_pr.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        _insert_before_first(
            tbl_pr,
            table_indent,
            ("w:tblBorders", "w:tblLayout", "w:tblLook"),
        )
    table_indent.set(qn("w:type"), "dxa")
    table_indent.set(qn("w:w"), "0")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge, color, size in (
        ("top", PRIMARY, "8"),
        ("left", PRIMARY, "8"),
        ("bottom", PRIMARY, "8"),
        ("right", PRIMARY, "8"),
        ("insideH", "C8D2D6", "4"),
        ("insideV", "C8D2D6", "4"),
    ):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)
        borders.append(element)
    _insert_before_first(tbl_pr, borders, ("w:tblLayout", "w:tblCellMar", "w:tblLook"))

    table_grid = table._tbl.tblGrid
    for column in list(table_grid.findall(qn("w:gridCol"))):
        table_grid.remove(column)
    for width in column_widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        table_grid.append(column)
    _set_table_cell_widths(table, column_widths_dxa)

    if table.rows:
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:tblHeader")) is None:
            tr_pr.append(OxmlElement("w:tblHeader"))

    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for cell in row.cells:
            if row_index == 0:
                _set_cell_shading(cell, LIGHT_HEADER)
            _set_cell_margins(cell, top=90, start=110, bottom=90, end=110)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = None
                paragraph.paragraph_format.space_after = Pt(1.5)
                paragraph.paragraph_format.line_spacing = 1.15


def _set_table_cell_widths(table, column_widths_dxa: tuple[int, ...]) -> None:
    for row_index, table_row in enumerate(table._tbl.findall(qn("w:tr")), start=1):
        row_properties = table_row.find(qn("w:trPr"))
        grid_before = (
            row_properties.find(qn("w:gridBefore"))
            if row_properties is not None
            else None
        )
        grid_after = (
            row_properties.find(qn("w:gridAfter"))
            if row_properties is not None
            else None
        )
        grid_index = int(grid_before.get(qn("w:val"))) if grid_before is not None else 0
        for cell_index, table_cell in enumerate(
            table_row.findall(qn("w:tc")), start=1
        ):
            cell_properties = table_cell.get_or_add_tcPr()
            grid_span = cell_properties.find(qn("w:gridSpan"))
            span = int(grid_span.get(qn("w:val"))) if grid_span is not None else 1
            if span < 1 or grid_index + span > len(column_widths_dxa):
                raise ValueError(
                    f"table row {row_index} cell {cell_index} has invalid grid span"
                )
            cell_width = cell_properties.get_or_add_tcW()
            cell_width.set(qn("w:type"), "dxa")
            cell_width.set(
                qn("w:w"), str(sum(column_widths_dxa[grid_index : grid_index + span]))
            )
            grid_index += span
        if grid_after is not None:
            grid_index += int(grid_after.get(qn("w:val")))
        if grid_index != len(column_widths_dxa):
            raise ValueError(f"table row {row_index} does not cover the table grid")


def shade_paragraph(paragraph, fill: str, *, border_color: str | None = None) -> None:
    """Apply a restrained background and optional left rule to a paragraph."""

    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    if border_color:
        borders = p_pr.find(qn("w:pBdr"))
        if borders is None:
            borders = OxmlElement("w:pBdr")
            _insert_before_first(
                p_pr,
                borders,
                ("w:shd", "w:tabs", "w:spacing", "w:ind", "w:jc"),
            )
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "6")
        left.set(qn("w:color"), border_color)
        borders.append(left)


def _configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    _configure_paragraph_style(
        normal,
        latin=LATIN_FONT,
        east_asia=CJK_SERIF_FONT,
        size=10.5,
        color="222222",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        line_spacing=1.45,
        after=4,
        first_indent=21,
    )

    _configure_paragraph_style(
        styles["Title"],
        latin=LATIN_FONT,
        east_asia=CJK_SANS_FONT,
        size=20,
        color=DARK,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.2,
        before=0,
        after=18,
    )
    _configure_paragraph_style(
        styles["Subtitle"],
        latin=LATIN_FONT,
        east_asia=CJK_SANS_FONT,
        size=12,
        color=BLUE_GRAY,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.25,
        after=8,
    )
    _configure_paragraph_style(
        styles["Heading 1"],
        latin=LATIN_FONT,
        east_asia=CJK_SANS_FONT,
        size=15,
        color=PRIMARY,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.2,
        before=12,
        after=6,
        keep_with_next=True,
    )
    _configure_paragraph_style(
        styles["Heading 2"],
        latin=LATIN_FONT,
        east_asia=CJK_SANS_FONT,
        size=13,
        color=PRIMARY,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.2,
        before=12,
        after=6,
        keep_with_next=True,
    )
    _configure_paragraph_style(
        styles["Heading 3"],
        latin=LATIN_FONT,
        east_asia=CJK_SANS_FONT,
        size=11.5,
        color=BLUE_GRAY,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.2,
        before=9,
        after=4,
        keep_with_next=True,
    )
    chapter_title = _get_or_add_style(styles, "Chapter Title")
    _configure_paragraph_style(
        chapter_title,
        latin=LATIN_FONT,
        east_asia=CJK_SANS_FONT,
        size=20,
        color=DARK,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.2,
        before=8,
        after=14,
        keep_with_next=True,
    )
    chapter_p_pr = chapter_title.element.get_or_add_pPr()
    outline = chapter_p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        chapter_p_pr.append(outline)
    outline.set(qn("w:val"), "0")
    _configure_paragraph_style(
        styles["Caption"],
        latin=LATIN_FONT,
        east_asia=CJK_SANS_FONT,
        size=9,
        color=BLUE_GRAY,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.15,
        before=3,
        after=6,
        keep_with_next=True,
    )

    code = _get_or_add_style(styles, "Code")
    _configure_paragraph_style(
        code,
        latin=MONO_FONT,
        east_asia=CJK_SANS_FONT,
        size=8.5,
        color=DARK,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.1,
        before=3,
        after=5,
    )
    shade_paragraph(_StyleParagraphAdapter(code), LIGHT_GRAY)

    callout = _get_or_add_style(styles, "Callout")
    _configure_paragraph_style(
        callout,
        latin=LATIN_FONT,
        east_asia=CJK_SERIF_FONT,
        size=10,
        color=DARK,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.3,
        before=4,
        after=5,
    )


class _StyleParagraphAdapter:
    """Expose a style's pPr through the small interface used by shading."""

    def __init__(self, style):
        self._p = style.element


def _configure_paragraph_style(
    style,
    *,
    latin: str,
    east_asia: str,
    size: float,
    color: str,
    bold: bool = False,
    alignment=None,
    line_spacing: float | None = None,
    before: float = 0,
    after: float = 0,
    first_indent: float | None = None,
    keep_with_next: bool = False,
    page_break_before: bool = False,
) -> None:
    font = style.font
    font.name = latin
    font.size = Pt(size)
    font.bold = bold
    font.color.rgb = RGBColor.from_string(color)
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for attribute, value in (
        ("w:ascii", latin),
        ("w:hAnsi", latin),
        ("w:eastAsia", east_asia),
        ("w:cs", latin),
    ):
        r_fonts.set(qn(attribute), value)
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "en-US")
    lang.set(qn("w:eastAsia"), "zh-CN")

    paragraph = style.paragraph_format
    paragraph.alignment = alignment
    paragraph.space_before = Pt(before)
    paragraph.space_after = Pt(after)
    paragraph.keep_with_next = keep_with_next
    paragraph.page_break_before = page_break_before
    paragraph.first_line_indent = Pt(first_indent) if first_indent else None
    if line_spacing is not None:
        paragraph.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph.line_spacing = line_spacing


def _set_document_defaults(document: Document) -> None:
    styles_element = document.styles.element
    doc_defaults = styles_element.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles_element.insert(0, doc_defaults)
    r_pr_default = doc_defaults.find(qn("w:rPrDefault"))
    if r_pr_default is None:
        r_pr_default = OxmlElement("w:rPrDefault")
        doc_defaults.append(r_pr_default)
    r_pr = r_pr_default.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        r_pr_default.append(r_pr)
    fonts = r_pr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.append(fonts)
    for attribute, value in (
        ("w:ascii", LATIN_FONT),
        ("w:hAnsi", LATIN_FONT),
        ("w:eastAsia", CJK_SERIF_FONT),
        ("w:cs", LATIN_FONT),
    ):
        fonts.set(qn(attribute), value)


def _get_or_add_style(styles, name: str):
    try:
        return styles[name]
    except KeyError:
        return styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def _format_run(run, *, size: float, color: str, east_asia: str) -> None:
    run.font.name = LATIN_FONT
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    r_fonts = run._r.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), LATIN_FONT)
    r_fonts.set(qn("w:hAnsi"), LATIN_FONT)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:cs"), LATIN_FONT)


def _clear_story(story) -> None:
    paragraph = story.paragraphs[0]
    paragraph.clear()
    for extra in list(story.paragraphs[1:]):
        extra._element.getparent().remove(extra._element)


def _set_page_numbering(section, page_format: str, page_start: int | None) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        _insert_before_first(
            sect_pr,
            pg_num,
            (
                "w:cols",
                "w:formProt",
                "w:vAlign",
                "w:titlePg",
                "w:textDirection",
                "w:bidi",
                "w:docGrid",
            ),
        )
    pg_num.set(qn("w:fmt"), page_format)
    if page_start is None:
        pg_num.attrib.pop(qn("w:start"), None)
    else:
        pg_num.set(qn("w:start"), str(page_start))


def _set_paragraph_bottom_border(paragraph, color: str, *, size: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        _insert_before_first(
            p_pr,
            borders,
            ("w:shd", "w:tabs", "w:spacing", "w:ind", "w:jc"),
        )
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def _set_cell_margins(cell, *, top: int, start: int, bottom: int, end: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        _insert_before_first(
            tc_pr,
            margins,
            ("w:textDirection", "w:tcFitText", "w:vAlign", "w:hideMark"),
        )
    for edge, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        _insert_before_first(
            tc_pr,
            shading,
            ("w:noWrap", "w:tcMar", "w:textDirection", "w:vAlign"),
        )
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)


def _create_list_numbering(document: Document) -> NumberingDefinition:
    numbering = document.part.numbering_part.element
    bullet_abstract_id = _next_numeric_attribute(
        numbering, "w:abstractNum", "w:abstractNumId"
    )
    decimal_abstract_id = bullet_abstract_id + 1
    _insert_abstract_numbering(
        numbering, _build_abstract_numbering(bullet_abstract_id, kind="bullet")
    )
    _insert_abstract_numbering(
        numbering, _build_abstract_numbering(decimal_abstract_id, kind="decimal")
    )
    bullet_num_id = add_numbering_instance(document, bullet_abstract_id)
    return NumberingDefinition(
        bullet_num_id=bullet_num_id,
        decimal_abstract_id=decimal_abstract_id,
    )


def _build_abstract_numbering(abstract_id: int, *, kind: str):
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)

    bullet_marks = ("•", "–", "▪")
    for level in range(9):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), kind)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(
            qn("w:val"),
            bullet_marks[level % len(bullet_marks)]
            if kind == "bullet"
            else f"%{level + 1}.",
        )
        justification = OxmlElement("w:lvlJc")
        justification.set(qn("w:val"), "left")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(420 + level * 360))
        tabs.append(tab)
        indent = OxmlElement("w:ind")
        indent.set(qn("w:left"), str(420 + level * 360))
        indent.set(qn("w:hanging"), "240")
        p_pr.extend((tabs, indent))
        lvl.extend((start, num_fmt, lvl_text, justification, p_pr))
        if kind == "bullet":
            r_pr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), CJK_SANS_FONT)
            fonts.set(qn("w:hAnsi"), CJK_SANS_FONT)
            r_pr.append(fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
    return abstract


def _insert_abstract_numbering(numbering, abstract) -> None:
    """Keep abstract definitions before numbering instances per OOXML order."""

    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(numbering.index(first_num), abstract)


def _next_numeric_attribute(parent, child_tag: str, attribute: str) -> int:
    values = [
        int(child.get(qn(attribute)))
        for child in parent.findall(qn(child_tag))
        if child.get(qn(attribute), "").isdigit()
    ]
    return max(values, default=0) + 1


def _insert_before_first(parent, child, following_tags: tuple[str, ...]) -> None:
    following = {qn(tag) for tag in following_tags}
    for existing in parent:
        if existing.tag in following:
            parent.insert(parent.index(existing), child)
            return
    parent.append(child)
