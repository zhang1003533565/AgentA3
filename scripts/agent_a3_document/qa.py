"""Artifact audits for the AgentA3 project document deliverables."""

from __future__ import annotations

import argparse
from collections import Counter
from dataclasses import dataclass
import ipaddress
import json
from pathlib import Path
import re
from typing import Any
from xml.etree import ElementTree
from zipfile import BadZipFile, ZipFile

import pdfplumber
from docx import Document
from pypdf import PdfReader


@dataclass
class AuditReport:
    """Collected QA findings and machine-readable metrics."""

    errors: list[str]
    warnings: list[str]
    metrics: dict[str, Any]


MIN_PAGE_COUNT = 70
MAX_PAGE_COUNT = 78
MIN_RASTER_DIMENSION_PX = 120

REQUIRED_CHAPTERS = (
    "第一章 项目概述",
    "第二章 需求分析",
    "第三章 总体设计",
    "第四章 详细功能设计",
    "第五章 核心技术设计",
    "第六章 数据与接口设计",
    "第七章 页面与交互设计",
    "第八章 测试与验证",
    "第九章 安装部署与使用",
    "第十章 项目总结与附录",
)

_BANNED_CLAIMS = (
    "所有智能体回答均经过 RAG",
    "30 个智能体同时自主协商",
    "考试结果已经自动更新七维画像",
    "已完成完整的动态学习路径管理",
    "所有资源包均包含视频或完整 PPTX",
    "一条 Compose 命令能够部署全部系统",
    "全部测试通过",
    "已通过等保",
    "已完成全部许可证审计",
)
_PLACEHOLDER_PATTERNS = (
    ("TBD", re.compile(r"(?<![A-Za-z0-9_])TBD(?![A-Za-z0-9_])", re.I)),
    ("TODO", re.compile(r"(?<![A-Za-z0-9_])TODO(?![A-Za-z0-9_])", re.I)),
    ("XXX", re.compile(r"(?<![A-Za-z0-9_])XXX(?![A-Za-z0-9_])", re.I)),
    ("X-X", re.compile(r"(?<![A-Za-z0-9_])X\s*-\s*X(?![A-Za-z0-9_])", re.I)),
    ("N 条样本", re.compile(r"N\s*条样本", re.I)),
    ("待填写", re.compile(r"待填写")),
    ("学校名称", re.compile(r"学校名称")),
    ("成员姓名", re.compile(r"成员姓名")),
    ("template marker", re.compile(r"\{\{[^{}\n]+\}\}|\[\[[^\[\]\n]+\]\]")),
)
_CREDENTIAL_RE = re.compile(
    r"(?im)(?P<key>password|passwd|pwd|secret|token|api[_-]?key|"
    r"access[_-]?key|client[_-]?secret|密码|密钥|令牌)"
    r"\s*[:=]\s*[`\"']?(?P<value>[^\s`\"',;]{4,})"
)
_AWS_ACCESS_KEY_RE = re.compile(r"(?<![A-Z0-9])AKIA[0-9A-Z]{16}(?![A-Z0-9])")
_PRIVATE_KEY_RE = re.compile(r"-----?BEGIN(?: [A-Z]+)? PRIVATE KEY-----?", re.I)
_IPV4_RE = re.compile(r"(?<![\d.])(?:\d{1,3}\.){3}\d{1,3}(?![\d.])")
_NEGATED_CLAIM_RE = re.compile(
    r"(?:不(?:宣称|表述为|认为|表示|意味着|等于|是)?|"
    r"未(?:宣称|表述为|实现|完成)?|并非|不能|不得)\s*$"
)
_ABSOLUTE_PATH_PATTERNS = (
    ("POSIX local path", re.compile(r"(?<![\w.])/(?:Users|home|private|tmp|var/folders)/[^\s\"'<>)]*")),
    ("Windows local path", re.compile(r"(?<![A-Za-z0-9_])[A-Za-z]:[\\/][^\s\"'<>)]*")),
    ("UNC local path", re.compile(r"(?<!\\)\\\\[^\s\\/]+[\\/][^\s\"'<>)]*")),
    ("file URI", re.compile(r"file:/+(?:Users|home|private|tmp|var|[A-Za-z]:)[^\s\"'<>)]*", re.I)),
)
_CAPTION_RE = re.compile(
    r"(?m)^\s*(?P<kind>[图表])\s*(?P<chapter>\d+)\s*[-－—–]\s*"
    r"(?P<sequence>\d+)(?=\s|$)"
)
_RELATIONSHIPS_NS = "http://schemas.openxmlformats.org/package/2006/relationships"


def _empty_metrics() -> dict[str, int]:
    return {
        "banned_claim_findings": 0,
        "placeholder_findings": 0,
        "sensitive_value_findings": 0,
        "absolute_path_findings": 0,
        "duplicate_caption_numbers": 0,
    }


def _audit_text(text: str) -> AuditReport:
    errors: list[str] = []
    metrics = _empty_metrics()

    for claim in _BANNED_CLAIMS:
        count = sum(
            1
            for match in re.finditer(re.escape(claim), text)
            if not _NEGATED_CLAIM_RE.search(text[max(0, match.start() - 24) : match.start()])
        )
        if not count:
            continue
        metrics["banned_claim_findings"] += count
        errors.append(f"[banned-claim] {claim!r} appears {count} time(s)")

    for label, pattern in _PLACEHOLDER_PATTERNS:
        count = len(pattern.findall(text))
        if not count:
            continue
        metrics["placeholder_findings"] += count
        errors.append(f"[placeholder] {label!r} appears {count} time(s)")

    sensitive_findings: list[str] = []
    for match in _CREDENTIAL_RE.finditer(text):
        value = match.group("value").strip().rstrip(".)]")
        lowered = value.casefold()
        if (
            set(value) <= {"*", "x", "X", "•"}
            or lowered in {"redacted", "masked", "removed", "example", "示例", "已脱敏"}
            or value.startswith("${")
        ):
            continue
        sensitive_findings.append(f"credential assignment for {match.group('key')}")
    sensitive_findings.extend(
        "AWS access-key-like value" for _ in _AWS_ACCESS_KEY_RE.finditer(text)
    )
    sensitive_findings.extend(
        "private-key material" for _ in _PRIVATE_KEY_RE.finditer(text)
    )
    for match in _IPV4_RE.finditer(text):
        try:
            address = ipaddress.ip_address(match.group(0))
        except ValueError:
            continue
        if address.is_global:
            sensitive_findings.append("public IPv4 address")
    metrics["sensitive_value_findings"] = len(sensitive_findings)
    errors.extend(
        f"[sensitive-value] {finding}" for finding in sensitive_findings
    )

    absolute_findings: list[str] = []
    for label, pattern in _ABSOLUTE_PATH_PATTERNS:
        absolute_findings.extend(label for _ in pattern.finditer(text))
    metrics["absolute_path_findings"] = len(absolute_findings)
    errors.extend(f"[absolute-path] {finding}" for finding in absolute_findings)

    return AuditReport(errors=errors, warnings=[], metrics=metrics)


def _caption_number(match: re.Match[str]) -> str:
    return f"{match.group('kind')}{int(match.group('chapter'))}-{int(match.group('sequence'))}"


def _caption_metrics(captions: list[str]) -> tuple[list[str], dict[str, int]]:
    counts = Counter(captions)
    duplicates = sorted(number for number, count in counts.items() if count > 1)
    errors = [
        f"[duplicate-caption] {number} appears {counts[number]} time(s)"
        for number in duplicates
    ]
    return errors, {
        "caption_count": len(captions),
        "duplicate_caption_numbers": len(duplicates),
    }


def _missing_chapters(*texts: str) -> list[str]:
    searchable = "\n".join(texts)
    return [chapter for chapter in REQUIRED_CHAPTERS if chapter not in searchable]


def _merge_text_and_caption_reports(
    text: str, captions: list[str], *, extra_metrics: dict[str, Any]
) -> AuditReport:
    report = _audit_text(text)
    caption_errors, caption_metrics = _caption_metrics(captions)
    report.errors.extend(caption_errors)
    report.metrics.update(caption_metrics)
    report.metrics.update(extra_metrics)
    return report


def audit_source(path: Path) -> AuditReport:
    path = Path(path)
    if path.is_file():
        files = [path]
    elif path.is_dir():
        files = sorted(candidate for candidate in path.rglob("*.md") if candidate.is_file())
    else:
        return AuditReport(
            errors=[f"[missing-artifact] source path does not exist: {path}"],
            warnings=[],
            metrics={**_empty_metrics(), "source_files": 0},
        )

    parts: list[str] = []
    for source in files:
        try:
            parts.append(source.read_text(encoding="utf-8"))
        except (OSError, UnicodeError) as exc:
            return AuditReport(
                errors=[f"[unreadable-artifact] cannot read source {source}: {exc}"],
                warnings=[],
                metrics={**_empty_metrics(), "source_files": len(files)},
            )
    text = "\n".join(parts)
    captions = [_caption_number(match) for match in _CAPTION_RE.finditer(text)]
    report = _merge_text_and_caption_reports(
        text,
        captions,
        extra_metrics={
            "source_files": len(files),
            "source_characters": len(text),
        },
    )
    missing = _missing_chapters(text)
    report.metrics["missing_chapters"] = missing
    report.errors.extend(
        f"[missing-chapter] {chapter}" for chapter in missing
    )
    return report


def audit_docx(path: Path) -> AuditReport:
    path = Path(path)
    if not path.is_file():
        return AuditReport(
            errors=[f"[missing-artifact] DOCX does not exist: {path}"],
            warnings=[],
            metrics={**_empty_metrics(), "docx_text_blocks": 0},
        )
    try:
        document = Document(path)
    except (OSError, ValueError, BadZipFile) as exc:
        return AuditReport(
            errors=[f"[unreadable-artifact] cannot parse DOCX {path}: {exc}"],
            warnings=[],
            metrics={**_empty_metrics(), "docx_text_blocks": 0},
        )

    text_blocks: list[str] = []
    captions: list[str] = []

    def add_paragraphs(paragraphs) -> None:
        for paragraph in paragraphs:
            value = paragraph.text.strip()
            if value:
                text_blocks.append(value)
            if getattr(paragraph.style, "name", "") == "Caption":
                match = _CAPTION_RE.match(value)
                if match:
                    captions.append(_caption_number(match))

    def add_tables(tables) -> None:
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    add_paragraphs(cell.paragraphs)
                    add_tables(cell.tables)

    add_paragraphs(document.paragraphs)
    add_tables(document.tables)
    for section in document.sections:
        for container in (section.header, section.footer):
            add_paragraphs(container.paragraphs)
            add_tables(container.tables)

    relationship_targets: list[str] = []
    try:
        with ZipFile(path) as package:
            for name in package.namelist():
                if not name.endswith(".rels"):
                    continue
                root = ElementTree.fromstring(package.read(name))
                relationship_targets.extend(
                    relationship.get("Target", "")
                    for relationship in root.findall(
                        f"{{{_RELATIONSHIPS_NS}}}Relationship"
                    )
                )
    except (BadZipFile, ElementTree.ParseError, OSError) as exc:
        return AuditReport(
            errors=[f"[unreadable-artifact] cannot inspect DOCX relationships: {exc}"],
            warnings=[],
            metrics={**_empty_metrics(), "docx_text_blocks": len(text_blocks)},
        )

    text = "\n".join((*text_blocks, *relationship_targets))
    report = _merge_text_and_caption_reports(
        text,
        captions,
        extra_metrics={
            "docx_text_blocks": len(text_blocks),
            "docx_relationships": len(relationship_targets),
        },
    )
    missing = _missing_chapters("\n".join(text_blocks))
    report.metrics["missing_chapters"] = missing
    report.errors.extend(
        f"[missing-chapter] {chapter}" for chapter in missing
    )
    return report


def audit_pdf(path: Path, *, allow_page_count_warning: bool = False) -> AuditReport:
    path = Path(path)
    if not path.is_file():
        return AuditReport(
            errors=[f"[missing-artifact] PDF does not exist: {path}"],
            warnings=[],
            metrics={**_empty_metrics(), "page_count": 0},
        )
    try:
        reader = PdfReader(path)
        page_count = len(reader.pages)
        outline_titles = list(_iter_outline_titles(reader.outline))
        with pdfplumber.open(path) as pdf:
            page_texts: list[str] = []
            blank_pages: list[int] = []
            captions: list[str] = []
            small_rasters: list[dict[str, int]] = []
            for page_number, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                page_texts.append(text)
                images = page.images
                if not text.strip() and not images:
                    blank_pages.append(page_number)
                for image in images:
                    source_size = image.get("srcsize")
                    if not source_size or len(source_size) != 2:
                        continue
                    width_px, height_px = map(int, source_size)
                    if min(width_px, height_px) < MIN_RASTER_DIMENSION_PX:
                        small_rasters.append(
                            {
                                "page": page_number,
                                "width_px": width_px,
                                "height_px": height_px,
                            }
                        )
                for line in page.extract_text_lines(layout=True, return_chars=False):
                    value = line.get("text", "").strip()
                    match = _CAPTION_RE.match(value)
                    if not match:
                        continue
                    line_center = (float(line["x0"]) + float(line["x1"])) / 2
                    if abs(line_center - float(page.width) / 2) <= float(page.width) * 0.05:
                        captions.append(_caption_number(match))
    except Exception as exc:
        return AuditReport(
            errors=[f"[unreadable-artifact] cannot parse PDF {path}: {exc}"],
            warnings=[],
            metrics={**_empty_metrics(), "page_count": 0},
        )

    combined_text = "\n".join(page_texts)
    report = _merge_text_and_caption_reports(
        combined_text,
        captions,
        extra_metrics={
            "page_count": page_count,
            "blank_pages": blank_pages,
            "small_raster_figures": len(small_rasters),
            "small_raster_details": small_rasters,
            "pdf_text_characters": len(combined_text),
            "outline_titles": len(outline_titles),
        },
    )
    if not MIN_PAGE_COUNT <= page_count <= MAX_PAGE_COUNT:
        finding = (
            f"[page-count] expected {MIN_PAGE_COUNT}-{MAX_PAGE_COUNT} pages, got {page_count}"
        )
        if allow_page_count_warning:
            report.warnings.append(finding)
        else:
            report.errors.append(finding)
    report.errors.extend(
        f"[blank-page] page {page_number} contains no text or raster content"
        for page_number in blank_pages
    )
    report.errors.extend(
        "[small-raster] page {page}: source raster is {width_px}x{height_px}px; "
        "minimum dimension is {minimum}px".format(
            **detail, minimum=MIN_RASTER_DIMENSION_PX
        )
        for detail in small_rasters
    )
    missing = _missing_chapters(combined_text, "\n".join(outline_titles))
    report.metrics["missing_chapters"] = missing
    report.errors.extend(
        f"[missing-chapter] {chapter}" for chapter in missing
    )

    cjk_chars = re.findall(r"[\u3400-\u9fff]", combined_text)
    if len(cjk_chars) >= 100:
        _, dominant_count = Counter(cjk_chars).most_common(1)[0]
        if dominant_count / len(cjk_chars) >= 0.3:
            report.warnings.append(
                "[text-extraction] PDF CJK extraction is highly repetitive; "
                "chapter validation used PDF outline titles as an additional source"
            )
    return report


def _iter_outline_titles(items):
    for item in items or ():
        if isinstance(item, list):
            yield from _iter_outline_titles(item)
            continue
        title = getattr(item, "title", None)
        if title is None and hasattr(item, "get"):
            title = item.get("/Title")
        if title:
            yield str(title)


def create_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Audit AgentA3 source, DOCX, and rendered PDF artifacts."
    )
    parser.add_argument(
        "--source",
        type=Path,
        default=Path("docs/project-document"),
        help="Source Markdown file or directory (default: docs/project-document).",
    )
    parser.add_argument("--docx", type=Path, required=True)
    parser.add_argument("--pdf", type=Path, required=True)
    parser.add_argument(
        "--allow-page-count-warning",
        action="store_true",
        help=(
            "Report an out-of-range PDF page count as a warning during approved "
            "wrap-up; all other audit findings remain errors."
        ),
    )
    return parser


def main(argv: list[str] | None = None) -> int:
    args = create_parser().parse_args(argv)
    reports = (
        ("source", audit_source(args.source)),
        ("docx", audit_docx(args.docx)),
        (
            "pdf",
            audit_pdf(
                args.pdf,
                allow_page_count_warning=args.allow_page_count_warning,
            ),
        ),
    )
    for label, report in reports:
        print(
            f"{label}: errors={len(report.errors)} warnings={len(report.warnings)} "
            f"metrics={json.dumps(report.metrics, ensure_ascii=False, sort_keys=True)}"
        )
        for error in report.errors:
            print(f"ERROR {label} {error}")
        for warning in report.warnings:
            print(f"WARNING {label} {warning}")
    return 1 if any(report.errors for _, report in reports) else 0


if __name__ == "__main__":
    raise SystemExit(main())
