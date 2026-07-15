from pathlib import Path, PurePosixPath
from hashlib import sha256
import re
from shutil import copy2
from tempfile import TemporaryDirectory
import unittest
from urllib.parse import unquote
from xml.etree import ElementTree
from zipfile import ZipFile

from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

from scripts.agent_a3_document.source_loader import (
    EvidenceRow,
    extract_evidence_rows,
    load_source,
    parse_markdown,
)


_RELATIONSHIPS_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
_WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_WINDOWS_ABSOLUTE_RE = re.compile(r"^[A-Za-z]:[\\/]")


def _relationship_violations(parts: dict[str, bytes]) -> list[tuple[str, str, str]]:
    violations: list[tuple[str, str, str]] = []
    for part_name, data in parts.items():
        if not part_name.endswith(".rels"):
            continue
        root = ElementTree.fromstring(data)
        for relationship in root.findall(f"{{{_RELATIONSHIPS_NS}}}Relationship"):
            target = relationship.get("Target", "")
            target_mode = relationship.get("TargetMode", "")
            if target_mode.casefold() == "external":
                violations.append((part_name, target, "external relationship"))
            if _is_local_or_escaping_target(part_name, target):
                violations.append((part_name, target, "local or escaping target"))
    return violations


def _is_local_or_escaping_target(relationship_part: str, target: str) -> bool:
    decoded = unquote(target).strip()
    lowered = decoded.casefold()
    if (
        lowered.startswith("file:")
        or decoded.startswith("/")
        or decoded.startswith("\\\\")
        or _WINDOWS_ABSOLUTE_RE.match(decoded)
    ):
        return True

    normalized = decoded.replace("\\", "/")
    relationship_path = PurePosixPath(relationship_part)
    if relationship_part == "_rels/.rels":
        stack: list[str] = []
    elif relationship_path.parent.name == "_rels":
        stack = list(relationship_path.parent.parent.parts)
    else:
        stack = list(relationship_path.parent.parts)

    for part in PurePosixPath(normalized).parts:
        if part in {"", "."}:
            continue
        if part == "..":
            if not stack:
                return True
            stack.pop()
        else:
            stack.append(part)
    return False


def _write_sources(root: Path, first_body: str) -> None:
    for index in range(11):
        body = first_body if index == 0 else f"# 第{index + 1}章\n\n正文。\n"
        (root / f"{index:02d}-chapter.md").write_text(body, encoding="utf-8")


class SourceContractTest(unittest.TestCase):
    def _assert_rejects_figure_width(self, width_cm: str) -> None:
        body = (
            '<!-- FIGURE src="figure.png" caption="验证图" '
            f'width_cm="{width_cm}" -->\n'
        )
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            _write_sources(root, body)
            with self.assertRaisesRegex(
                ValueError,
                r"00-chapter\.md:1: figure width_cm must be a positive number",
            ):
                load_source(root)

    def test_loads_required_block_types_in_numeric_order(self):
        chapter_names = ["第一章 项目概述"] + [
            f"第{i}章 验证章节" for i in range(2, 11)
        ] + ["附录"]
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            for index, chapter in enumerate(chapter_names):
                body = (
                    f'<!-- SECTION_BREAK chapter="{chapter}" -->\n'
                    f"# {chapter}\n\n正文。\n\n"
                    '<!-- FIGURE src="figure.png" caption="验证图" '
                    'width_cm="15.5" -->\n\n'
                    "| 字段 | 内容 |\n|---|---|\n| 状态 | 已验证 |\n"
                )
                (root / f"{index:02d}-chapter.md").write_text(
                    body, encoding="utf-8"
                )
            blocks = load_source(root)
        kinds = {block.kind for block in blocks}
        self.assertTrue(
            {"heading", "paragraph", "table", "figure", "section_break"}
            <= kinds
        )
        chapters = [b.text for b in blocks if b.kind == "section_break"]
        self.assertEqual(chapters[0], "第一章 项目概述")
        self.assertEqual(chapters[-1], "附录")

    def test_loads_every_approved_block_kind(self):
        body = """<!-- AUTO_TOC -->
<!-- SECTION_BREAK chapter="第一章 项目概述" -->
# 第一章 项目概述

正文第一行。
正文第二行。

- 一级项目
  - 二级项目
1. 第一项
2. 第二项

```python
print("ok")
```

<!-- FIGURE src="images/overview.png" caption="总体图" width_cm="15.5" -->

| 字段 | 内容 |
|---|---|
| 状态 | 已验证 |

<!-- CALLOUT type="risk" title="当前限制" -->
第一段。

第二段。
<!-- END_CALLOUT -->

<!-- PAGE_BREAK -->
"""
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            _write_sources(root, body)
            blocks = load_source(root)

        self.assertEqual(
            {block.kind for block in blocks},
            {
                "heading",
                "paragraph",
                "bullet",
                "numbered",
                "table",
                "figure",
                "callout",
                "code",
                "toc",
                "page_break",
                "section_break",
            },
        )
        bullets = [block for block in blocks if block.kind == "bullet"]
        self.assertEqual(
            [(block.text, block.level) for block in bullets],
            [("一级项目", 1), ("二级项目", 2)],
        )
        code = next(block for block in blocks if block.kind == "code")
        self.assertEqual(code.text, 'print("ok")')
        self.assertEqual(code.attrs, {"language": "python"})
        callout = next(block for block in blocks if block.kind == "callout")
        self.assertEqual(callout.text, "第一段。\n\n第二段。")
        self.assertEqual(callout.attrs, {"type": "risk", "title": "当前限制"})

    def test_rejects_unknown_directive_with_source_location(self):
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            _write_sources(root, "<!-- UNKNOWN value=\"x\" -->\n")
            with self.assertRaisesRegex(
                ValueError, r"00-chapter\.md:1: unknown directive"
            ):
                load_source(root)

    def test_rejects_non_numeric_figure_width(self):
        self._assert_rejects_figure_width("wide")

    def test_rejects_zero_figure_width(self):
        self._assert_rejects_figure_width("0")

    def test_rejects_negative_figure_width(self):
        self._assert_rejects_figure_width("-1.5")

    def test_rejects_malformed_table_with_source_location(self):
        body = """# 第一章

| 字段 | 内容 |
|---|
| 状态 | 已验证 |
"""
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            _write_sources(root, body)
            with self.assertRaisesRegex(
                ValueError, r"00-chapter\.md:3: malformed table"
            ):
                load_source(root)

    def test_requires_exactly_eleven_chapter_sources(self):
        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            (root / "00-only.md").write_text("# 唯一章节\n", encoding="utf-8")
            with self.assertRaisesRegex(
                ValueError, r"expected 11 chapter sources, found 1"
            ):
                load_source(root)


class EvidenceIndexContractTest(unittest.TestCase):
    def test_extracts_rows_with_the_four_approved_statuses(self):
        statuses = ("implemented", "partial", "planned", "known-limit")
        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "evidence-index.md"
            data_rows = "\n".join(
                f"| EV-{index:03d} | 结论 {index} | {status} | "
                f"`path/to/evidence-{index}.md` | 成品表述 {index} |"
                for index, status in enumerate(statuses, start=1)
            )
            path.write_text(
                "# 工程证据索引\n\n"
                "| ID | Claim | Status | Evidence | Final wording |\n"
                "|---|---|---|---|---|\n"
                f"{data_rows}\n",
                encoding="utf-8",
            )

            rows = extract_evidence_rows(path)

        self.assertEqual(
            rows,
            [
                EvidenceRow(
                    id=f"EV-{index:03d}",
                    claim=f"结论 {index}",
                    status=status,
                    evidence=f"`path/to/evidence-{index}.md`",
                    final_wording=f"成品表述 {index}",
                )
                for index, status in enumerate(statuses, start=1)
            ],
        )

    def test_rejects_absolute_evidence_paths(self):
        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "evidence-index.md"
            path.write_text(
                "| ID | Claim | Status | Evidence | Final wording |\n"
                "|---|---|---|---|---|\n"
                "| EV-001 | 不安全证据 | implemented | `/tmp/secret.txt` | "
                "不得使用 |\n",
                encoding="utf-8",
            )
            with self.assertRaisesRegex(
                ValueError,
                r"evidence-index\.md:3: evidence paths must be repository-relative",
            ):
                extract_evidence_rows(path)

    def test_rejects_windows_drive_and_unc_evidence_paths(self):
        unsafe_paths = (
            "C:/Users/name/file.txt",
            r"C:\Users\name\file.txt",
            r"\\server\share\file.txt",
            "//server/share/file.txt",
        )
        for unsafe_path in unsafe_paths:
            with self.subTest(unsafe_path=unsafe_path), TemporaryDirectory() as tmp:
                path = Path(tmp) / "evidence-index.md"
                path.write_text(
                    "| ID | Claim | Status | Evidence | Final wording |\n"
                    "|---|---|---|---|---|\n"
                    f"| EV-001 | 不安全证据 | implemented | `{unsafe_path}` | "
                    "不得使用 |\n",
                    encoding="utf-8",
                )
                with self.assertRaisesRegex(
                    ValueError,
                    r"evidence-index\.md:3: evidence paths must be "
                    r"repository-relative",
                ):
                    extract_evidence_rows(path)


class SourceContentTest(unittest.TestCase):
    SOURCE_DIR = Path("docs/project-document/source")
    SOURCE_NAMES = (
        "00-frontmatter.md",
        "01-project-overview.md",
        "02-requirements.md",
        "03-overall-design.md",
    )

    def _read_source(self, name: str) -> str:
        path = self.SOURCE_DIR / name
        self.assertTrue(path.is_file(), f"missing reviewed source: {path}")
        return path.read_text(encoding="utf-8")

    def _combined_source(self) -> str:
        return "\n".join(self._read_source(name) for name in self.SOURCE_NAMES)

    def _requirement_row(self, requirement_id: str) -> tuple[str, ...]:
        text = self._read_source("02-requirements.md")
        prefix = f"| {requirement_id} |"
        line = next(
            (candidate for candidate in text.splitlines() if candidate.startswith(prefix)),
            None,
        )
        self.assertIsNotNone(line, f"missing requirement row: {requirement_id}")
        return tuple(cell.strip() for cell in line.strip().strip("|").split("|"))

    def test_frontmatter_uses_exact_titles_and_document_governance(self):
        text = self._read_source("00-frontmatter.md")
        required_text = (
            "AgentA3 智慧校园个性化学习多智能体系统——需求、设计与开发说明书",
            "第十五届中国软件杯大赛——A组赛题",
            "A3-基于大模型的个性化资源生成与学习多智能体系统开发",
            "AgentA3 项目组",
            "V1.0",
            "2026-07-15",
            "修订记录",
            "摘要",
            "文档范围",
        )
        for value in required_text:
            with self.subTest(value=value):
                self.assertIn(value, text)
        self.assertEqual(text.count("<!-- AUTO_TOC -->"), 1)

    def test_sources_use_only_supported_directives_and_section_breaks(self):
        for name in self.SOURCE_NAMES:
            with self.subTest(name=name):
                text = self._read_source(name)
                blocks = parse_markdown(text, source_name=name)
                self.assertTrue(blocks)
                self.assertEqual(
                    sum(block.kind == "section_break" for block in blocks), 1
                )

    def test_requirements_have_exact_unique_ids_and_traceability_columns(self):
        text = self._read_source("02-requirements.md")
        functional_ids = re.findall(r"(?<![A-Z0-9-])FR-\d{3}(?![A-Z0-9-])", text)
        nonfunctional_ids = re.findall(
            r"(?<![A-Z0-9-])NFR-\d{3}(?![A-Z0-9-])", text
        )
        self.assertEqual(
            functional_ids, [f"FR-{index:03d}" for index in range(1, 25)]
        )
        self.assertEqual(
            nonfunctional_ids,
            [f"NFR-{index:03d}" for index in range(1, 11)],
        )
        traceability_header = (
            "| ID | 参与者（actor） | 触发条件（trigger） | "
            "系统行为（system behavior） | 验收证据（acceptance evidence） | "
            "当前状态（current status） |"
        )
        self.assertEqual(text.count(traceability_header), 2)

    def test_requirements_cover_roles_pain_points_and_capability_matrix(self):
        text = self._read_source("02-requirements.md")
        for role in ("ADMIN", "TEACHER", "STUDENT", "MERCHANT"):
            with self.subTest(role=role):
                self.assertIn(role, text)
        pain_point_ids = re.findall(r"(?<![A-Z0-9-])BP-\d{3}(?![A-Z0-9-])", text)
        self.assertEqual(pain_point_ids, [f"BP-{index:03d}" for index in range(1, 7)])
        capability_header = (
            "| 能力维度 | 通用聊天机器人 | 传统教学平台 | "
            "基础知识库问答 | AgentA3 |"
        )
        self.assertEqual(text.count(capability_header), 1)

    def test_current_entry_and_admin_only_permissions_are_explicit(self):
        requirements = self._read_source("02-requirements.md")
        design = self._read_source("03-overall-design.md")

        for role_row in (
            "| ADMIN | 管理员 | Web |",
            "| TEACHER | 教师 | App |",
            "| STUDENT | 学生 | App |",
            "| MERCHANT | 商户 | Web |",
        ):
            with self.subTest(role_row=role_row):
                self.assertIn(role_row, requirements)

        entry_fact = (
            "当前 App 登录入口允许 STUDENT 与 TEACHER；"
            "当前 Web 登录入口允许 ADMIN 与 MERCHANT。"
        )
        self.assertIn(entry_fact, requirements)
        self.assertIn(entry_fact, design)

        expected_rows = {
            "FR-002": ("学生、教师", "已实现（implemented）"),
            "FR-012": ("管理员", "已实现（implemented）"),
            "FR-013": (
                "管理员（当前）；学生、教师（目标）",
                "部分实现（partial）",
            ),
            "FR-014": (
                "管理员（当前）；学生、教师（目标）",
                "部分实现（partial）",
            ),
            "FR-018": (
                "管理员（当前）；教师（目标）",
                "部分实现（partial）",
            ),
            "FR-019": (
                "管理员（当前）；教师（目标）",
                "部分实现（partial）",
            ),
            "FR-020": (
                "管理员、商户（当前可进入）；教师（目标）",
                "当前限制（known-limit）",
            ),
            "FR-021": (
                "管理员、商户（当前可进入）；教师（目标）",
                "当前限制（known-limit）",
            ),
            "FR-024": (
                "管理员、商户（当前可登录）；教师（目标）",
                "当前限制（known-limit）",
            ),
        }
        for requirement_id, (actor, status_prefix) in expected_rows.items():
            with self.subTest(requirement_id=requirement_id):
                row = self._requirement_row(requirement_id)
                self.assertEqual(row[1], actor)
                self.assertTrue(row[5].startswith(status_prefix), row[5])

        for requirement_id in ("FR-012", "FR-013", "FR-014"):
            with self.subTest(requirement_id=requirement_id):
                self.assertIn("ADMIN", " ".join(self._requirement_row(requirement_id)))
        for requirement_id in ("FR-018", "FR-019"):
            with self.subTest(requirement_id=requirement_id):
                self.assertIn("ADMIN", " ".join(self._requirement_row(requirement_id)))

    def test_second_review_boundaries_are_explicit(self):
        requirements = self._read_source("02-requirements.md")
        design = self._read_source("03-overall-design.md")

        paper_auth_fact = (
            "当前 Web 路由与导航未按角色过滤；试卷创建、列表、详情与预览"
            "只校验登录，只有发布和取消发布执行 ADMIN 校验。"
        )
        self.assertIn(paper_auth_fact, requirements)
        self.assertIn(paper_auth_fact, design)
        merchant_row = next(
            line
            for line in requirements.splitlines()
            if line.startswith("| MERCHANT |")
        )
        self.assertIn("当前授权缺口", merchant_row)
        self.assertIn("未来加固", merchant_row)

        retrieval_fact = (
            "MaxKB 仅执行 hit-test 检索；Java 提取引用并组装 grounded context；"
            "系统 LLM/agent 生成最终回答，Java 再组合 citations 响应。"
        )
        self.assertIn(retrieval_fact, requirements)
        self.assertIn(retrieval_fact, design)

        python_boundary_terms = (
            "Python 会反向调用 Java，并转发原始 Authorization",
            "Redis 不可用时退化为进程内存",
            "`AI_PYTHON_BASE_URL`",
            "`JAVA_BACKEND_BASE_URL`",
            "`REDIS_URL`",
            "`AI_INTERNAL_TOKEN`",
            "`X-AI-Internal-Token`",
        )
        for term in python_boundary_terms:
            with self.subTest(term=term):
                self.assertIn(term, design)

        self.assertNotIn("MaxKB 检索与回答", design)
        self.assertIn(
            "MaxKB hit-test 检索；系统 LLM/agent 生成最终回答",
            design,
        )
        self.assertIn(
            "Java、Python 与 Redis 地址均可由环境覆盖，Compose 使用内部 DNS 名称。",
            design,
        )

        for stale_fact in (
            "Java base URL 当前固定为 `http://localhost:8080`",
            "Python 记忆服务直接尝试 `redis://localhost:6379/0`",
            "本地默认地址尚未环境配置",
        ):
            with self.subTest(stale_fact=stale_fact):
                self.assertNotIn(stale_fact, design)

        requirements_cancellation_fact = (
            "服务端取消接口属于设计要求；当前客户端 AbortController 仅中止本地 fetch，"
            "不能证明存在 REST 取消端点。"
        )
        design_cancellation_fact = (
            "普通助手的服务端取消端点没有现有证据，客户端 AbortController 只中止本地 fetch。"
        )
        self.assertIn(requirements_cancellation_fact, requirements)
        self.assertIn(design_cancellation_fact, design)

    def test_overall_design_names_components_flows_and_deployment_boundary(self):
        text = self._read_source("03-overall-design.md")
        required_terms = (
            "uni-app",
            "React/Vite",
            "Spring Boot",
            "FastAPI",
            "MySQL",
            "Redis",
            "MaxKB",
            "讯飞（Xfyun）",
            "REST",
            "SSE",
            "WebSocket",
            "已纳入五服务 Compose/CI 静态契约",
            "deploy/compose.submission.yml",
            "不作为独立公网 API",
        )
        for term in required_terms:
            with self.subTest(term=term):
                self.assertIn(term, text)

    def test_overall_design_compaction_preserves_architecture_contracts(self):
        text = self._read_source("03-overall-design.md")
        for heading in (
            "3.1 设计原则",
            "3.2 系统上下文",
            "3.3 四部分总体架构",
            "3.4 Java 与 Python 职责分界",
            "3.5 数据与外部依赖",
            "3.6 通信协议与主要数据流",
            "3.7 逻辑模块划分",
            "3.8 当前实现",
            "3.9 部分实现",
            "3.10 部署边界",
            "3.11 安全边界",
            "3.12 后续规划",
        ):
            with self.subTest(heading=heading):
                self.assertIn(heading, text)
        for fact in (
            "当前 App 登录入口允许 STUDENT 与 TEACHER",
            "当前 Web 登录入口允许 ADMIN 与 MERCHANT",
            "试卷创建、列表、详情与预览只校验登录",
            "MaxKB 仅执行 hit-test 检索",
            "系统 LLM/agent 生成最终回答",
            "Python 会反向调用 Java，并转发原始 Authorization",
            "`AI_PYTHON_BASE_URL`",
            "`JAVA_BACKEND_BASE_URL`",
            "`REDIS_URL`",
            "`AI_INTERNAL_TOKEN`",
            "Redis 不可用时退化为进程内存",
            "客户端 AbortController 只中止本地 fetch",
            "已纳入五服务 Compose/CI 静态契约",
            "`/internal/*`",
        ):
            with self.subTest(fact=fact):
                self.assertIn(fact, text)

    def test_sources_separate_requirements_goals_facts_and_roadmap(self):
        text = self._combined_source()
        for heading in (
            "赛题要求",
            "设计目标",
            "当前实现",
            "部分实现",
            "后续规划",
        ):
            with self.subTest(heading=heading):
                self.assertIn(heading, text)
        for status in ("implemented", "partial", "planned", "known-limit"):
            with self.subTest(status=status):
                self.assertIn(status, text)

    def test_sources_exclude_banned_claims_placeholders_secrets_and_local_paths(self):
        text = self._combined_source()
        banned_claims = (
            "所有智能体回答均经过 RAG",
            "30 个智能体同时自主协商",
            "考试结果已经自动更新七维画像",
            "已完成完整的动态学习路径管理",
            "所有资源包均包含视频或完整 PPTX",
            "一条 Compose 命令能够部署全部系统",
            "全部测试通过",
            "已通过等保",
        )
        forbidden_tokens = (
            "TBD",
            "TODO",
            "XXX",
            "X-X",
            "N 条样本",
            "待填写",
            "学校名称",
            "成员姓名",
            "/Users/",
            "/home/",
            "C:\\Users\\",
            "BEGIN PRIVATE KEY",
            "AKIA",
        )
        for value in (*banned_claims, *forbidden_tokens):
            with self.subTest(value=value):
                self.assertNotIn(value, text)


class DetailedFunctionsAndCoreTechnologiesSourceTest(unittest.TestCase):
    SOURCE_DIR = Path("docs/project-document/source")
    FUNCTION_SOURCE = "04-detailed-functions.md"
    TECHNOLOGY_SOURCE = "05-core-technologies.md"

    def _read_source(self, name: str) -> str:
        path = self.SOURCE_DIR / name
        self.assertTrue(path.is_file(), f"missing reviewed source: {path}")
        return path.read_text(encoding="utf-8")

    @staticmethod
    def _numbered_sections(text: str, prefix: str) -> list[str]:
        pattern = rf"(?m)^## [^\n]*\b{prefix}-\d{{2}}\b[^\n]*$"
        matches = list(re.finditer(pattern, text))
        return [
            text[match.start() : matches[index + 1].start()]
            if index + 1 < len(matches)
            else text[match.start() :]
            for index, match in enumerate(matches)
        ]

    def test_chapter_sources_exist_and_use_supported_markdown(self):
        expected_chapters = {
            self.FUNCTION_SOURCE: "第四章 详细功能设计",
            self.TECHNOLOGY_SOURCE: "第五章 核心技术设计",
        }
        for name, chapter in expected_chapters.items():
            with self.subTest(name=name):
                text = self._read_source(name)
                blocks = parse_markdown(text, source_name=name)
                self.assertTrue(blocks)
                self.assertEqual(
                    [block.text for block in blocks if block.kind == "section_break"],
                    [chapter],
                )

    def test_detailed_functions_have_exact_unique_ids_and_atomic_contracts(self):
        text = self._read_source(self.FUNCTION_SOURCE)
        function_ids = re.findall(
            r"(?m)^## [^\n]*\b(FUNC-\d{2})\b[^\n]*$", text
        )
        self.assertEqual(function_ids, [f"FUNC-{index:02d}" for index in range(1, 12)])
        self.assertEqual(
            re.findall(r"(?<![A-Z0-9-])FUNC-\d{2}(?![A-Z0-9-])", text),
            function_ids,
        )

        required_labels = (
            "**目标**：",
            "**参与者**：",
            "**触发条件**：",
            "**前置条件**：",
            "**主流程**：",
            "**替代流程**：",
            "**异常流程**：",
            "**后置条件**：",
            "**业务规则**：",
            "**涉及接口**：",
            "**数据实体**：",
            "**验收条件**：",
            "**测试编号**：",
        )
        sections = self._numbered_sections(text, "FUNC")
        self.assertEqual(len(sections), 11)
        for function_id, section in zip(function_ids, sections):
            for label in required_labels:
                with self.subTest(function_id=function_id, label=label):
                    self.assertIn(label, section)
            self.assertRegex(section, rf"FT-{function_id}-[A-Z0-9-]+")

    def test_core_technologies_have_evidence_limits_and_figure_directives(self):
        text = self._read_source(self.TECHNOLOGY_SOURCE)
        technology_ids = re.findall(
            r"(?m)^## [^\n]*\b(TECH-\d{2})\b[^\n]*$", text
        )
        self.assertEqual(technology_ids, [f"TECH-{index:02d}" for index in range(1, 9)])
        self.assertEqual(
            re.findall(r"(?<![A-Z0-9-])TECH-\d{2}(?![A-Z0-9-])", text),
            technology_ids,
        )

        evidence_ids = {
            row.id
            for row in extract_evidence_rows(
                Path("docs/project-document/evidence-index.md")
            )
        }
        sections = self._numbered_sections(text, "TECH")
        self.assertEqual(len(sections), 8)
        for technology_id, section in zip(technology_ids, sections):
            with self.subTest(technology_id=technology_id):
                self.assertEqual(
                    section.count(
                        '<!-- CALLOUT type="evidence" title="证据映射" -->'
                    ),
                    1,
                )
                self.assertEqual(section.count("<!-- END_CALLOUT -->"), 1)
                self.assertIn("**失败与降级**：", section)
                self.assertIn("**已知限制**：", section)
                self.assertRegex(
                    section,
                    r'<!-- FIGURE src="images/tech-[^"]+\.png" '
                    r'caption="[^"]+" width_cm="15\.5" -->',
                )
                mapped_ids = set(re.findall(r"\bEV-\d{3}\b", section))
                self.assertTrue(mapped_ids, f"{technology_id} has no evidence mapping")
                self.assertTrue(mapped_ids <= evidence_ids)

    def test_current_fact_boundaries_are_explicit(self):
        functions = self._read_source(self.FUNCTION_SOURCE)
        technologies = self._read_source(self.TECHNOLOGY_SOURCE)
        combined = f"{functions}\n{technologies}"

        required_facts = (
            "当前 App 登录入口允许 STUDENT 与 TEACHER；当前 Web 登录入口允许 ADMIN 与 MERCHANT。",
            "试卷创建、列表、详情与预览只校验登录，只有发布和取消发布执行 ADMIN 校验",
            "MaxKB 仅执行 hit-test 检索",
            "系统 LLM/agent 生成最终回答",
            "Python 会反向调用 Java，并转发原始 Authorization",
            "Redis 不可用时退化为进程内存",
            "30 个专业智能体实现包",
            "candidate",
            "applied",
            "受支持的客观题交卷后按题目知识点",
            "同一提交不重复更新掌握度、证据或路径",
            "主观题不自动评分，也不进入客观题反馈映射",
            "讯飞（Xfyun）实时 ASR WebSocket",
            "顺序执行",
            "不构成零幻觉证明",
            "七个题目智能体实现包",
            "当前 Web 与考试业务链路只暴露五类题型",
            "一次性预览证明",
            "客观题自动评分",
            "不对简答题执行自动评分",
        )
        for fact in required_facts:
            with self.subTest(fact=fact):
                self.assertIn(fact, combined)

        banned_claims = (
            "30 个智能体同时自主协商",
            "会议具备完整 RTC",
            "会议具备 TTS",
            "资源信封保证零幻觉",
            "已实现简答题自动评分",
            "考试结果已经自动更新七维画像",
            "所有智能体回答均经过 RAG",
        )
        for claim in banned_claims:
            with self.subTest(claim=claim):
                self.assertNotIn(claim, combined)


class RemainingProjectDocumentSourceTest(unittest.TestCase):
    SOURCE_DIR = Path("docs/project-document/source")
    SOURCE_NAMES = (
        "00-frontmatter.md",
        "01-project-overview.md",
        "02-requirements.md",
        "03-overall-design.md",
        "04-detailed-functions.md",
        "05-core-technologies.md",
        "06-data-interfaces.md",
        "07-ui-interaction.md",
        "08-testing.md",
        "09-deployment-usage.md",
        "10-summary-appendices.md",
    )
    REMAINING_SOURCES = SOURCE_NAMES[6:]

    def _read_remaining(self, name: str) -> str:
        path = self.SOURCE_DIR / name
        self.assertTrue(path.is_file(), f"missing reviewed source: {path}")
        return path.read_text(encoding="utf-8")

    def _combined_remaining(self) -> str:
        return "\n".join(self._read_remaining(name) for name in self.REMAINING_SOURCES)

    def test_exactly_eleven_sources_load_with_required_chapters(self):
        self.assertEqual(
            tuple(path.name for path in sorted(self.SOURCE_DIR.glob("[0-9][0-9]-*.md"))),
            self.SOURCE_NAMES,
        )
        blocks = load_source(self.SOURCE_DIR)
        self.assertEqual(
            [block.text for block in blocks if block.kind == "section_break"],
            [
                "前置材料",
                "第一章 项目概述",
                "第二章 需求分析",
                "第三章 总体设计",
                "第四章 详细功能设计",
                "第五章 核心技术设计",
                "第六章 数据与接口设计",
                "第七章 页面与交互设计",
                "第八章 测试与验证",
                "第九章 安装部署与使用",
                "第十章 项目总结与附录",
            ],
        )

    def test_remaining_sources_use_only_supported_directives(self):
        for name in self.REMAINING_SOURCES:
            with self.subTest(name=name):
                text = self._read_remaining(name)
                blocks = parse_markdown(text, source_name=name)
                self.assertTrue(blocks)
                self.assertEqual(
                    sum(block.kind == "section_break" for block in blocks), 1
                )

    def test_data_and_interface_scale_and_boundaries_are_explicit(self):
        text = self._read_remaining("06-data-interfaces.md")
        required = (
            "65 个 JPA 实体",
            "64 个唯一映射表",
            "逻辑 ID 关系不能作为物理外键已经存在的证明",
            "54 个控制器",
            "404 个映射注解",
            "MaxKB 仅执行 hit-test 检索",
            "系统 LLM/agent 生成最终回答",
            "Python 会反向调用 Java，并转发原始 Authorization",
            "当前 Web 路由与导航未按角色过滤",
        )
        for value in required:
            with self.subTest(value=value):
                self.assertIn(value, text)
        for trace_id in ("DATA-01", "DATA-02", "API-01", "API-02", "API-03"):
            with self.subTest(trace_id=trace_id):
                self.assertIn(trace_id, text)

    def test_ui_uses_only_the_three_approved_design_images(self):
        text = self._read_remaining("07-ui-interaction.md")
        figures = re.findall(
            r'<!-- FIGURE src="([^"]+)" caption="([^"]+)" width_cm="15\.5" -->',
            text,
        )
        self.assertEqual(len(figures), 3)
        self.assertEqual(
            {src for src, _ in figures},
            {
                "images/profile-radar-design.png",
                "images/ai-conversation-design.png",
                "images/meeting-home-design.png",
            },
        )
        self.assertTrue(all("界面设计稿" in caption for _, caption in figures))
        self.assertNotIn("docs/images/app-run.png", text)
        for value in ("SSE", "资源卡片", "下载", "失败恢复", "ASR", "自动保存"):
            with self.subTest(value=value):
                self.assertIn(value, text)

    def test_ui_compaction_preserves_interaction_and_authorization_contracts(self):
        text = self._read_remaining("07-ui-interaction.md")
        for heading in (
            "7.1 设计目标与页面边界",
            "7.2 页面清单与角色入口",
            "7.3 移动首页与导航",
            "7.4 七维画像雷达交互",
            "7.5 AI 对话与 SSE 状态",
            "7.6 六资源工作流、卡片与单项重试",
            "7.7 会议首页、实时 ASR 与会后结果",
            "7.8 题库、组卷与预览交互",
            "7.9 在线考试与自动保存",
            "7.10 Web 治理页面与配置交互",
            "7.11 通用反馈、恢复与可访问性",
        ):
            with self.subTest(heading=heading):
                self.assertIn(heading, text)
        for fact in (
            "当前 App 登录入口允许 STUDENT 与 TEACHER",
            "当前 Web 登录入口允许 ADMIN 与 MERCHANT",
            "随机预览、创建、列表、详情、预览与下载只校验登录",
            "当前 AbortController 只中止本地 fetch",
            "MaxKB 仅执行 hit-test 检索",
            "candidate",
            "applied",
            "learningUpdate",
            "简答题不得显示为已自动判分",
            "grounding",
            "integrity",
            "五类题型",
            "七个题目智能体",
            "一次性预览证明",
            "客观题",
            "简答题",
            "partial",
            "final",
        ):
            with self.subTest(fact=fact):
                self.assertIn(fact, text)

    def test_compacted_appendix_preserves_direct_dependency_bom_and_indices(self):
        text = self._read_remaining("10-summary-appendices.md")
        dependency_versions = (
            ("Spring Boot Parent", "4.0.3"),
            ("jjwt-api", "0.12.5"),
            ("jjwt-impl", "0.12.5"),
            ("jjwt-jackson", "0.12.5"),
            ("springdoc-openapi-starter-webmvc-ui", "2.8.4"),
            ("aliyun-sdk-oss", "3.17.4"),
            ("cos_api", "5.6.263"),
            ("playwright", "1.49.0"),
            ("poi-ooxml", "5.5.1"),
            ("pdfbox", "3.0.5"),
            ("@ant-design/v5-patch-for-react-19", "^1.0.3"),
            ("antd", "^5.29.3"),
            ("axios", "^1.13.6"),
            ("dayjs", "^1.11.20"),
            ("echarts", "^6.0.0"),
            ("react", "^19.2.0"),
            ("react-dom", "^19.2.0"),
            ("react-markdown", "^10.1.0"),
            ("react-router-dom", "^7.13.1"),
            ("vite", "^7.3.1"),
            ("compressorjs", "^1.3.0"),
            ("fastapi", "0.115.12"),
            ("uvicorn", "0.30.6"),
            ("langchain-openai", "0.2.14"),
            ("langchain-core", "0.3.29"),
            ("redis", "5.2.1"),
            ("pydantic", "2.10.6"),
            ("PyMuPDF", "1.25.5"),
            ("python-docx", "1.1.2"),
            ("python-pptx", "1.0.2"),
            ("pymilvus", "2.4.9"),
            ("setuptools", "80.9.0"),
            ("marshmallow", "3.26.1"),
        )
        for dependency, version in dependency_versions:
            with self.subTest(dependency=dependency):
                self.assertIn(dependency, text)
                self.assertIn(version, text)
        self.assertEqual(
            re.findall(r"(?m)^\[(\d+)\] ", text),
            [str(index) for index in range(1, 19)],
        )
        self.assertIn("`pdf2docx` 未被仓库代码使用", text)
        self.assertNotIn("| pdf2docx |", text)
        for evidence_id in ("EV-001", "EV-011", "EV-022", "EV-032", "EV-048"):
            with self.subTest(evidence_id=evidence_id):
                self.assertIn(evidence_id, text)

    def test_paper_auth_gap_includes_random_preview_and_download(self):
        disclosure = "随机预览、创建、列表、详情、预览与下载只校验登录"
        for name in (
            "06-data-interfaces.md",
            "07-ui-interaction.md",
            "09-deployment-usage.md",
            "10-summary-appendices.md",
        ):
            with self.subTest(name=name):
                self.assertIn(disclosure, self._read_remaining(name))

    def test_generated_figure_names_follow_asset_manifest(self):
        data = self._read_remaining("06-data-interfaces.md")
        deployment = self._read_remaining("09-deployment-usage.md")
        self.assertIn('src="images/core-entity-relations.png"', data)
        self.assertIn('src="images/deployment-boundary.png"', deployment)
        self.assertNotIn('src="images/data-entity-relations.png"', data)
        self.assertNotIn('src="images/deployment-topology.png"', deployment)

    def test_testing_deployment_and_disclosures_match_current_evidence(self):
        testing = self._read_remaining("08-testing.md")
        deployment = self._read_remaining("09-deployment-usage.md")
        appendices = self._read_remaining("10-summary-appendices.md")

        for result in (
            "389 tests / 0 failures / 0 errors / 1 skipped",
            "Python 264 passed / 5 warnings",
            "AppWeb Node 46/46",
            "AppFrontend 93/93",
            "不代表最终整合 SHA",
        ):
            with self.subTest(result=result):
                self.assertIn(result, testing)
        self.assertEqual(
            re.findall(r"(?m)^\| (TC-\d{2}) \|", testing),
            [f"TC-{index:02d}" for index in range(1, 12)],
        )
        for gap in ("性能测试", "真机测试", "渗透测试", "生产端到端测试"):
            self.assertIn(gap, testing)

        for value in (
            "MySQL 8",
            "Redis 7",
            "Java 21",
            "React/Vite/Nginx",
            "HBuilderX",
            "FastAPI",
            "deploy/compose.submission.yml",
            "Python `ai-server` 与 Web 五个服务",
            "`AI_PYTHON_BASE_URL`",
            "`JAVA_BACKEND_BASE_URL`",
            "`REDIS_URL`",
            "`AI_INTERNAL_TOKEN`",
            "未确认镜像实际构建、`up -d`",
        ):
            with self.subTest(value=value):
                self.assertIn(value, deployment)

        disclosures = (
            "当前业务密码仍采用明文比较",
            "JWT 与第三方活动凭据已从跟踪配置值中移除",
            "数据库开发配置保留默认口令",
            "Git 历史 secret/PII 扫描、轮换",
            "WebSocket Origin 当前过于宽松",
            "仅列直接依赖及其版本",
            "许可证复核尚未完成",
            "许可证 P0 阻断",
            "OpenAI Codex（Desktop/本地工作区代理）",
            "oh-my-codex（OMX）协作层",
            "最终结果由项目成员审阅并承担责任",
            "GB/T 7714",
            "仓库相对路径",
        )
        for value in disclosures:
            with self.subTest(value=value):
                self.assertIn(value, appendices)

    def test_deployment_compaction_preserves_commands_roles_and_failure_boundaries(self):
        text = self._read_remaining("09-deployment-usage.md")
        for heading in (
            "9.1 部署范围与当前边界",
            "9.2 环境清单",
            "9.3 配置与敏感信息管理",
            "9.4 部署拓扑与启动顺序",
            "9.5 MySQL 8 与 Redis 7",
            "9.6 Java 21 服务",
            "9.7 Python FastAPI AI 服务",
            "9.8 React/Vite/Nginx Web 端",
            "9.9 uni-app 与 HBuilderX",
            "9.10 外部服务接入",
            "9.11 首次启动验收",
            "9.12 用户操作概要",
            "9.13 运维与故障排查",
        ):
            with self.subTest(heading=heading):
                self.assertIn(heading, text)
        for command in (
            "config --services",
            "config --quiet",
            "build",
            "up -d",
            "ps",
            "mvn -version",
            "mvn spring-boot:run",
            "./start-ai-server.sh",
            "npm install",
            "npm run build",
        ):
            with self.subTest(command=command):
                self.assertIn(command, text)
        for fact in (
            "deploy/compose.submission.yml",
            "Python `ai-server` 与 Web 五个服务",
            "退化为进程内存",
            "`AI_PYTHON_BASE_URL`",
            "`JAVA_BACKEND_BASE_URL`",
            "`REDIS_URL`",
            "`AI_INTERNAL_TOKEN`",
            "未确认镜像实际构建、`up -d`",
            "随机预览、创建、列表、详情、预览与下载只校验登录",
            "STUDENT",
            "TEACHER",
            "ADMIN",
            "MERCHANT",
        ):
            with self.subTest(fact=fact):
                self.assertIn(fact, text)

    def test_remaining_sources_exclude_placeholders_secrets_and_absolute_paths(self):
        text = self._combined_remaining()
        forbidden = (
            "TBD",
            "TODO",
            "XXX",
            "待填写",
            "/Users/",
            "/home/",
            "C:\\Users\\",
            "BEGIN PRIVATE KEY",
            "AKIA",
            "全部测试通过",
            "一条 Compose 命令能够部署全部系统",
            "已完成全部许可证审计",
        )
        for value in forbidden:
            with self.subTest(value=value):
                self.assertNotIn(value, text)


class GeneratedDiagramContractTest(unittest.TestCase):
    EXPECTED_DIAGRAMS = {
        "system-context.png",
        "four-part-architecture.png",
        "deployment-boundary.png",
        "agent-capability-groups.png",
        "profile-evidence-flow.png",
        "meeting-asr-loop.png",
        "maxkb-grounding-flow.png",
        "resource-envelope.png",
        "question-paper-exam-loop.png",
        "core-entity-relations.png",
        "requirements-trace.png",
    }
    APPROVED_COLORS = {
        (15, 118, 110),
        (18, 59, 69),
        (79, 107, 122),
        (183, 121, 31),
        (242, 244, 245),
    }

    def test_builds_exact_deterministic_diagram_manifest(self):
        from scripts.agent_a3_document.diagrams import build_all_diagrams

        with TemporaryDirectory() as tmp:
            output_dir = Path(tmp)
            generated = build_all_diagrams(output_dir)

            self.assertEqual({path.name for path in generated}, self.EXPECTED_DIAGRAMS)
            self.assertEqual(
                {path.name for path in output_dir.iterdir() if path.is_file()},
                self.EXPECTED_DIAGRAMS,
            )
            for path in generated:
                with self.subTest(path=path.name), Image.open(path) as image:
                    self.assertEqual(image.format, "PNG")
                    self.assertGreaterEqual(image.width, 1600)
                    self.assertGreaterEqual(image.height, 850)
                    self.assertIsNotNone(image.getbbox())
                    self.assertLessEqual(
                        set(image.convert("RGB").getdata()), self.APPROVED_COLORS
                    )

    def test_reviewed_semantic_manifests_are_explicit(self):
        from scripts.agent_a3_document.diagrams import (
            MEETING_RETURN_FLOW,
            QUESTION_PAPER_FINAL_CARD,
            QUESTION_PAPER_GROUP,
            SYSTEM_CONTEXT_EXTERNAL_CONNECTIONS,
            TRACE_ROWS,
            TRACE_STATUS_COLORS,
        )

        self.assertEqual(
            SYSTEM_CONTEXT_EXTERNAL_CONNECTIONS,
            (("Java", "讯飞 Xfyun ASR"), ("Python", "模型提供方")),
        )
        self.assertEqual(
            MEETING_RETURN_FLOW,
            (
                "Xfyun partial / final → Java handler",
                "Java → 客户端广播",
                "Java → MeetingRecord 持久化 final",
            ),
        )
        self.assertEqual(
            [row.test_id for row in TRACE_ROWS],
            [f"TC-{index:02d}" for index in range(1, 12)],
        )
        self.assertEqual(
            [row.status for row in TRACE_ROWS],
            [
                "known-limit",
                "partial",
                "implemented",
                "implemented",
                "partial",
                "partial",
                "partial",
                "known-limit",
                "partial",
                "known-limit",
                "partial",
            ],
        )
        self.assertGreater(len({row.status for row in TRACE_ROWS}), 1)
        self.assertTrue({row.status for row in TRACE_ROWS} <= TRACE_STATUS_COLORS.keys())
        self.assertLessEqual(QUESTION_PAPER_FINAL_CARD[2], QUESTION_PAPER_GROUP[2])

    def test_repeated_builds_are_byte_identical(self):
        from scripts.agent_a3_document.diagrams import build_all_diagrams

        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            first = build_all_diagrams(root / "first")
            second = build_all_diagrams(root / "second")

            self.assertEqual([path.name for path in first], [path.name for path in second])
            for first_path, second_path in zip(first, second):
                with self.subTest(path=first_path.name):
                    self.assertEqual(first_path.read_bytes(), second_path.read_bytes())


class ArtifactAuditInterfaceTest(unittest.TestCase):
    def test_qa_module_exposes_required_audit_interface(self):
        try:
            from scripts.agent_a3_document import qa
        except ImportError as exc:
            self.fail(f"missing Task 7 QA module: {exc}")

        self.assertTrue(callable(qa.audit_source))
        self.assertTrue(callable(qa.audit_docx))
        self.assertTrue(callable(qa.audit_pdf))
        self.assertTrue(callable(getattr(qa, "main", None)))
        report = qa.AuditReport(errors=[], warnings=[], metrics={})
        self.assertEqual(report.errors, [])
        self.assertEqual(report.warnings, [])
        self.assertEqual(report.metrics, {})


class ArtifactAuditBehaviorTest(unittest.TestCase):
    CHAPTERS = (
        "第一章 项目概述",
        "第二章 需求分析",
        "第三章 总体设计",
        "第四章 详细功能设计",
        "第五章 核心技术设计",
        "第六章 数据与接口设计",
        "第七章 页面与交互设计",
        "第八章 测试与验证",
        "第九章 安装部署与使用",
        "第十章 项目总结与附录",
    )

    def _assert_error_categories(self, report, expected):
        for category in expected:
            with self.subTest(category=category):
                self.assertTrue(
                    any(error.startswith(f"[{category}]") for error in report.errors),
                    f"missing {category!r} in {report.errors!r}",
                )

    def _write_pdf(self, path: Path, pages: list[list[str]], *, small_image=False):
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.pdfgen import canvas

        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        writer = canvas.Canvas(str(path), pagesize=A4)
        width, height = A4
        for page_index, lines in enumerate(pages):
            chapter = next((line for line in lines if line in self.CHAPTERS), None)
            if chapter is not None:
                bookmark = f"chapter-{page_index + 1}"
                writer.bookmarkPage(bookmark)
                writer.addOutlineEntry(chapter, bookmark, level=0)
            writer.setFont("STSong-Light", 10)
            y = height - 54
            for line in lines:
                if re.match(r"^[图表]\s*\d+\s*-\s*\d+", line):
                    writer.drawCentredString(width / 2, y, line)
                else:
                    writer.drawString(54, y, line)
                y -= 16
            if page_index == 0 and small_image:
                image_path = path.with_suffix(".png")
                Image.new("RGB", (16, 16), "#0F766E").save(image_path)
                writer.drawImage(str(image_path), 54, 80, width=144, height=144)
            writer.showPage()
        writer.save()

    def test_source_audit_rejects_claims_placeholders_credentials_paths_and_duplicate_captions(self):
        from scripts.agent_a3_document.qa import audit_source

        fixture = """# 第一章 项目概述
全部测试通过
TODO
api_key = sk-live-sensitive-value
/Users/example/work/AgentA3
图 1-1 第一张图
图 1-1 重复图
表 1-1 第一张表
表 1-1 重复表
"""
        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "source.md"
            path.write_text(fixture, encoding="utf-8")
            report = audit_source(path)

        self._assert_error_categories(
            report,
            {
                "banned-claim",
                "placeholder",
                "sensitive-value",
                "absolute-path",
                "duplicate-caption",
            },
        )
        self.assertEqual(report.metrics["banned_claim_findings"], 1)
        self.assertEqual(report.metrics["placeholder_findings"], 1)
        self.assertEqual(report.metrics["sensitive_value_findings"], 1)
        self.assertEqual(report.metrics["absolute_path_findings"], 1)
        self.assertEqual(report.metrics["duplicate_caption_numbers"], 2)

    def test_source_audit_allows_explicitly_negated_banned_claims(self):
        from scripts.agent_a3_document.qa import audit_source

        chapters = "\n".join(self.CHAPTERS)
        fixture = (
            f"{chapters}\n"
            "该说明不宣称所有智能体回答均经过 RAG。\n"
            "当前实现不表述为 30 个智能体同时自主协商。\n"
        )
        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "source.md"
            path.write_text(fixture, encoding="utf-8")
            report = audit_source(path)

        self.assertFalse(
            any(error.startswith("[banned-claim]") for error in report.errors),
            report.errors,
        )
        self.assertEqual(report.metrics["banned_claim_findings"], 0)

    def test_docx_audit_scans_body_tables_headers_footers_and_relationships(self):
        from scripts.agent_a3_document.qa import audit_docx

        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "unsafe.docx"
            document = Document()
            document.add_paragraph("全部测试通过")
            document.add_paragraph("TBD")
            table = document.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "password: hunter2"
            document.sections[0].header.paragraphs[0].text = "/home/example/AgentA3"
            document.sections[0].footer.paragraphs[0].text = "TODO"
            document.add_paragraph("图 2-1 第一张图", style="Caption")
            document.add_paragraph("图 2-1 重复图", style="Caption")
            document.save(path)

            report = audit_docx(path)

        self._assert_error_categories(
            report,
            {
                "banned-claim",
                "placeholder",
                "sensitive-value",
                "absolute-path",
                "duplicate-caption",
            },
        )
        self.assertGreaterEqual(report.metrics["docx_text_blocks"], 7)
        self.assertEqual(report.metrics["duplicate_caption_numbers"], 1)

    def test_pdf_audit_reports_content_layout_and_chapter_failures(self):
        from scripts.agent_a3_document.qa import audit_pdf

        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "unsafe.pdf"
            pages = [
                [
                    "第一章 项目概述",
                    "全部测试通过",
                    "XXX",
                    "token = sk-live-sensitive-value",
                    "C:\\Users\\example\\AgentA3",
                    "图 1-1 第一张图",
                    "图 1-1 重复图",
                ],
                [],
            ]
            self._write_pdf(path, pages, small_image=True)
            report = audit_pdf(path)

        self._assert_error_categories(
            report,
            {
                "page-count",
                "blank-page",
                "missing-chapter",
                "small-raster",
                "duplicate-caption",
                "banned-claim",
                "placeholder",
                "sensitive-value",
                "absolute-path",
            },
        )
        self.assertEqual(report.metrics["page_count"], 2)
        self.assertEqual(report.metrics["blank_pages"], [2])
        self.assertEqual(report.metrics["small_raster_figures"], 1)
        self.assertEqual(report.metrics["duplicate_caption_numbers"], 1)
        self.assertEqual(len(report.metrics["missing_chapters"]), 9)

    def test_pdf_audit_rejects_page_counts_below_and_above_bounds(self):
        from pypdf import PdfWriter

        from scripts.agent_a3_document.qa import audit_pdf

        with TemporaryDirectory() as tmp:
            for count in (104, 126):
                with self.subTest(page_count=count):
                    path = Path(tmp) / f"{count}.pdf"
                    writer = PdfWriter()
                    for _ in range(count):
                        writer.add_blank_page(width=595, height=842)
                    with path.open("wb") as output:
                        writer.write(output)
                    report = audit_pdf(path)
                    self.assertTrue(
                        any(
                            error.startswith("[page-count]")
                            for error in report.errors
                        )
                    )
                    self.assertEqual(report.metrics["page_count"], count)

    def test_pdf_audit_accepts_current_document_page_range(self):
        from pypdf import PdfWriter

        from scripts.agent_a3_document.qa import audit_pdf

        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "115.pdf"
            writer = PdfWriter()
            for _ in range(115):
                writer.add_blank_page(width=595, height=842)
            with path.open("wb") as output:
                writer.write(output)
            report = audit_pdf(path)

        self.assertFalse(
            any(error.startswith("[page-count]") for error in report.errors),
            report.errors,
        )

    def test_wrap_up_mode_demotes_only_page_count_to_warning(self):
        from scripts.agent_a3_document.qa import audit_pdf, main

        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "source"
            source.mkdir()
            chapter_text = "\n".join(self.CHAPTERS)
            (source / "chapters.md").write_text(chapter_text, encoding="utf-8")

            docx = root / "document.docx"
            document = Document()
            for chapter in self.CHAPTERS:
                document.add_paragraph(chapter)
            document.save(docx)

            pdf = root / "document.pdf"
            pages = [[*self.CHAPTERS, "正文内容"]]
            pages.extend([f"正文内容，第 {number} 页"] for number in range(2, 127))
            self._write_pdf(pdf, pages)

            report = audit_pdf(pdf, allow_page_count_warning=True)
            exit_code = main(
                [
                    "--source",
                    str(source),
                    "--docx",
                    str(docx),
                    "--pdf",
                    str(pdf),
                    "--allow-page-count-warning",
                ]
            )

        self.assertFalse(
            any(error.startswith("[page-count]") for error in report.errors),
            report.errors,
        )
        self.assertTrue(
            any(warning.startswith("[page-count]") for warning in report.warnings),
            report.warnings,
        )
        self.assertEqual(exit_code, 0)

    def test_qa_cli_returns_nonzero_when_any_artifact_has_errors(self):
        from scripts.agent_a3_document.qa import main

        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "source"
            source.mkdir()
            (source / "chapter.md").write_text("TODO\n", encoding="utf-8")
            docx = root / "unsafe.docx"
            document = Document()
            document.add_paragraph("TODO")
            document.save(docx)
            pdf = root / "unsafe.pdf"
            self._write_pdf(pdf, [["TODO"]])

            exit_code = main(
                [
                    "--source",
                    str(source),
                    "--docx",
                    str(docx),
                    "--pdf",
                    str(pdf),
                ]
            )

        self.assertNotEqual(exit_code, 0)


class ProjectDocumentDocxContractTest(unittest.TestCase):
    SOURCE_DIR = Path("docs/project-document/source")
    EVIDENCE_PATH = Path("docs/project-document/evidence-index.md")
    ASSET_DIR = Path("docs/project-document/assets/generated")
    DESIGN_IMAGES = (
        Path("docs/designs/profile-radar/profile-radar-design.png"),
        Path("docs/designs/ai-conversation/ai-conversation-design.png"),
        Path("docs/designs/meeting-home/meeting-home-design.png"),
    )
    CHAPTERS = (
        "前置材料",
        "第一章 项目概述",
        "第二章 需求分析",
        "第三章 总体设计",
        "第四章 详细功能设计",
        "第五章 核心技术设计",
        "第六章 数据与接口设计",
        "第七章 页面与交互设计",
        "第八章 测试与验证",
        "第九章 安装部署与使用",
        "第十章 项目总结与附录",
    )

    @classmethod
    def setUpClass(cls):
        from scripts.agent_a3_document.docx_builder import build_document

        cls._tmp = TemporaryDirectory()
        cls.output = Path(cls._tmp.name) / "AgentA3-project-document.docx"
        result = build_document(
            cls.SOURCE_DIR,
            cls.EVIDENCE_PATH,
            cls.ASSET_DIR,
            cls.output,
        )
        if result != cls.output:
            raise AssertionError(f"builder returned {result!r}, expected {cls.output!r}")
        if not cls.output.is_file():
            raise AssertionError(f"builder did not create {cls.output}")
        cls.document = Document(cls.output)
        with ZipFile(cls.output) as package:
            cls.parts = {name: package.read(name) for name in package.namelist()}

    @classmethod
    def tearDownClass(cls):
        if hasattr(cls, "_tmp"):
            cls._tmp.cleanup()

    def test_a4_geometry_sections_headers_and_page_number_starts(self):
        sections = self.document.sections
        self.assertGreaterEqual(len(sections), 12)
        for section in sections:
            self.assertEqual(section.orientation, WD_ORIENT.PORTRAIT)
            self.assertAlmostEqual(section.page_width.cm, 21.0, places=1)
            self.assertAlmostEqual(section.page_height.cm, 29.7, places=1)
            self.assertAlmostEqual(section.left_margin.cm, 2.2, places=1)
            self.assertAlmostEqual(section.right_margin.cm, 2.2, places=1)
            self.assertAlmostEqual(section.top_margin.cm, 2.3, places=1)
            self.assertAlmostEqual(section.bottom_margin.cm, 2.0, places=1)
            self.assertAlmostEqual(section.header_distance.cm, 1.25, places=1)
            self.assertAlmostEqual(section.footer_distance.cm, 1.25, places=1)

        self.assertFalse("".join(p.text for p in sections[0].header.paragraphs).strip())
        self.assertFalse("".join(p.text for p in sections[0].footer.paragraphs).strip())
        header_text = [
            " ".join(p.text for p in section.header.paragraphs).strip()
            for section in sections[1:]
        ]
        self.assertEqual(len(header_text), len(self.CHAPTERS))
        for actual, chapter in zip(header_text, self.CHAPTERS):
            self.assertIn("AgentA3", actual)
            self.assertIn(chapter, actual)

        front_num = sections[1]._sectPr.find(qn("w:pgNumType"))
        chapter_num = sections[2]._sectPr.find(qn("w:pgNumType"))
        self.assertIsNotNone(front_num)
        self.assertEqual(front_num.get(qn("w:fmt")), "upperRoman")
        self.assertEqual(front_num.get(qn("w:start")), "1")
        self.assertIsNotNone(chapter_num)
        self.assertEqual(chapter_num.get(qn("w:fmt")), "decimal")
        self.assertEqual(chapter_num.get(qn("w:start")), "1")

    def test_normal_style_uses_compact_readable_pagination_rhythm(self):
        paragraph_format = self.document.styles["Normal"].paragraph_format
        self.assertEqual(paragraph_format.line_spacing, 1.45)
        self.assertAlmostEqual(paragraph_format.space_after.pt, 4.0, places=1)

    def test_required_styles_have_explicit_latin_and_east_asian_fonts(self):
        required = (
            "Normal",
            "Title",
            "Subtitle",
            "Heading 1",
            "Heading 2",
            "Heading 3",
            "Caption",
            "Table Text",
            "Code",
            "Callout",
        )
        for name in required:
            with self.subTest(style=name):
                style = self.document.styles[name]
                fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
                self.assertTrue(fonts.get(qn("w:ascii")))
                self.assertTrue(fonts.get(qn("w:hAnsi")))
                self.assertTrue(fonts.get(qn("w:eastAsia")))

        normal = self.document.styles["Normal"]
        title = self.document.styles["Title"]
        heading_one = self.document.styles["Heading 1"]
        table_text = self.document.styles["Table Text"]
        self.assertEqual(normal.font.name, "Times New Roman")
        self.assertEqual(normal.font.size.pt, 10.5)
        self.assertEqual(table_text.font.size.pt, 9.0)
        self.assertEqual(title.font.size.pt, 20.0)
        self.assertEqual(heading_one.font.size.pt, 15.0)
        self.assertEqual(str(heading_one.font.color.rgb), "0F766E")

    def test_cjk_styles_and_runs_use_runtime_available_pingfang(self):
        xml_parts = [
            data.decode("utf-8")
            for name, data in self.parts.items()
            if name == "word/document.xml"
            or name == "word/styles.xml"
            or name == "word/numbering.xml"
            or name.startswith("word/header")
            or name.startswith("word/footer")
        ]
        combined_xml = "\n".join(xml_parts)
        self.assertNotIn("Source Han Serif SC", combined_xml)
        self.assertNotIn("Source Han Sans SC", combined_xml)
        self.assertIn("PingFang SC", combined_xml)

        for name in (
            "Normal",
            "Title",
            "Subtitle",
            "Heading 1",
            "Heading 2",
            "Heading 3",
            "Caption",
            "Table Text",
            "Code",
            "Callout",
        ):
            with self.subTest(style=name):
                fonts = (
                    self.document.styles[name]
                    ._element.get_or_add_rPr()
                    .get_or_add_rFonts()
                )
                self.assertEqual(fonts.get(qn("w:eastAsia")), "PingFang SC")

    def test_cover_has_no_references_and_each_footer_has_one_centered_page_field(self):
        word = {"w": _WORD_NS}
        document_root = ElementTree.fromstring(self.parts["word/document.xml"])
        section_properties = document_root.findall(".//w:sectPr", word)
        self.assertEqual(len(section_properties), 12)
        self.assertFalse(section_properties[0].findall("w:headerReference", word))
        self.assertFalse(section_properties[0].findall("w:footerReference", word))

        relationship_root = ElementTree.fromstring(
            self.parts["word/_rels/document.xml.rels"]
        )
        targets_by_id = {
            relationship.get("Id"): relationship.get("Target")
            for relationship in relationship_root.findall(
                f"{{{_RELATIONSHIPS_NS}}}Relationship"
            )
        }
        footer_targets: list[str] = []
        for section_index, section in enumerate(section_properties[1:], start=1):
            references = section.findall("w:footerReference", word)
            self.assertEqual(len(references), 1, f"section {section_index}")
            relationship_id = references[0].get(qn("r:id"))
            target = targets_by_id[relationship_id]
            footer_targets.append(target)

            footer_root = ElementTree.fromstring(self.parts[f"word/{target}"])
            page_paragraphs = [
                paragraph
                for paragraph in footer_root.findall("w:p", word)
                if [
                    node
                    for node in paragraph.findall(".//w:instrText", word)
                    if (node.text or "").strip() == "PAGE"
                ]
            ]
            self.assertEqual(len(page_paragraphs), 1, target)
            page_fields = [
                node
                for node in page_paragraphs[0].findall(".//w:instrText", word)
                if (node.text or "").strip() == "PAGE"
            ]
            self.assertEqual(len(page_fields), 1, target)
            justification = page_paragraphs[0].find("w:pPr/w:jc", word)
            self.assertIsNotNone(justification, target)
            self.assertEqual(justification.get(qn("w:val")), "center", target)

        self.assertEqual(len(set(footer_targets)), 11)

    def test_cover_title_breaks_at_the_document_scope_boundary(self):
        title = next(
            paragraph
            for paragraph in self.document.paragraphs
            if paragraph.style.name == "Title"
        )
        self.assertEqual(
            title.text,
            "AgentA3 智慧校园个性化学习多智能体系统\n"
            "——需求、设计与开发说明书",
        )
        self.assertEqual(len(title._p.findall(".//w:br", {"w": _WORD_NS})), 1)

    def test_every_table_has_consistent_fixed_dxa_geometry(self):
        word = {"w": _WORD_NS}
        usable_width_dxa = round(16.6 * 1440 / 2.54)
        document_root = ElementTree.fromstring(self.parts["word/document.xml"])
        tables = document_root.findall(".//w:tbl", word)
        source_table_count = sum(
            block.kind == "table" for block in load_source(self.SOURCE_DIR)
        )
        self.assertEqual(len(tables), source_table_count + 11)

        for table_index, table in enumerate(tables, start=1):
            table_width = table.find("w:tblPr/w:tblW", word)
            table_indent = table.find("w:tblPr/w:tblInd", word)
            table_layout = table.find("w:tblPr/w:tblLayout", word)
            self.assertIsNotNone(table_width, table_index)
            self.assertEqual(table_width.get(qn("w:type")), "dxa", table_index)
            self.assertEqual(
                int(table_width.get(qn("w:w"))), usable_width_dxa, table_index
            )
            self.assertIsNotNone(table_indent, table_index)
            self.assertEqual(table_indent.get(qn("w:type")), "dxa", table_index)
            self.assertEqual(table_indent.get(qn("w:w")), "0", table_index)
            self.assertIsNotNone(table_layout, table_index)
            self.assertEqual(table_layout.get(qn("w:type")), "fixed", table_index)

            grid_widths = [
                int(column.get(qn("w:w")))
                for column in table.findall("w:tblGrid/w:gridCol", word)
            ]
            self.assertTrue(grid_widths, table_index)
            self.assertEqual(sum(grid_widths), usable_width_dxa, table_index)
            for row_index, row in enumerate(table.findall("w:tr", word), start=1):
                row_properties = row.find("w:trPr", word)
                grid_before = (
                    row_properties.find("w:gridBefore", word)
                    if row_properties is not None
                    else None
                )
                grid_after = (
                    row_properties.find("w:gridAfter", word)
                    if row_properties is not None
                    else None
                )
                grid_index = int(grid_before.get(qn("w:val"))) if grid_before else 0
                for cell_index, cell in enumerate(row.findall("w:tc", word), start=1):
                    span_element = cell.find("w:tcPr/w:gridSpan", word)
                    span = int(span_element.get(qn("w:val"))) if span_element else 1
                    cell_width = cell.find("w:tcPr/w:tcW", word)
                    self.assertIsNotNone(
                        cell_width, f"table {table_index} row {row_index} cell {cell_index}"
                    )
                    self.assertEqual(cell_width.get(qn("w:type")), "dxa")
                    self.assertEqual(
                        int(cell_width.get(qn("w:w"))),
                        sum(grid_widths[grid_index : grid_index + span]),
                        f"table {table_index} row {row_index} cell {cell_index}",
                    )
                    grid_index += span
                grid_index += int(grid_after.get(qn("w:val"))) if grid_after else 0
                self.assertEqual(grid_index, len(grid_widths), table_index)

    def test_merged_cells_receive_the_sum_of_their_grid_columns(self):
        from scripts.agent_a3_document.styles import configure_table

        document = Document()
        table = document.add_table(rows=1, cols=3)
        table.cell(0, 0).merge(table.cell(0, 1))
        configure_table(table, (3000, 3000, 3411))
        word = {"w": _WORD_NS}
        grid_widths = [
            int(column.get(qn("w:w")))
            for column in table._tbl.findall("w:tblGrid/w:gridCol", word)
        ]
        cell_widths = [
            int(cell.find("w:tcPr/w:tcW", word).get(qn("w:w")))
            for cell in table._tbl.findall("w:tr/w:tc", word)
        ]
        self.assertEqual(grid_widths, [3000, 3000, 3411])
        self.assertEqual(cell_widths, [6000, 3411])

    def test_table_cells_use_compact_readable_vertical_rhythm(self):
        from scripts.agent_a3_document.styles import configure_table

        document = Document()
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "多行表格正文"
        configure_table(table, (9411,))

        margins = table.cell(0, 0)._tc.get_or_add_tcPr().find(qn("w:tcMar"))
        self.assertIsNotNone(margins)
        self.assertEqual(margins.find(qn("w:top")).get(qn("w:w")), "70")
        self.assertEqual(margins.find(qn("w:bottom")).get(qn("w:w")), "70")
        paragraph_format = table.cell(0, 0).paragraphs[0].paragraph_format
        self.assertEqual(paragraph_format.line_spacing, 1.10)
        self.assertAlmostEqual(paragraph_format.space_after.pt, 0.5, places=1)
        self.assertIsNone(table.rows[0]._tr.get_or_add_trPr().find(qn("w:cantSplit")))

    def test_document_tables_use_dedicated_nine_point_table_text_style(self):
        for table_index, table in enumerate(self.document.tables, start=1):
            for row_index, row in enumerate(table.rows, start=1):
                for cell_index, cell in enumerate(row.cells, start=1):
                    for paragraph in cell.paragraphs:
                        self.assertEqual(
                            paragraph.style.name,
                            "Table Text",
                            f"table {table_index} row {row_index} cell {cell_index}",
                        )

    def test_function_sections_are_materialized_as_compact_traceable_tables(self):
        function_tables = [
            table
            for table in self.document.tables
            if table.cell(0, 0).text.startswith("FUNC-")
        ]
        self.assertEqual(len(function_tables), 11)
        expected_labels = (
            "目标",
            "参与者",
            "触发条件",
            "前置条件",
            "主流程",
            "替代流程",
            "异常流程",
            "后置条件",
            "业务规则",
            "涉及接口",
            "数据实体",
            "验收条件",
            "测试编号",
        )
        for index, table in enumerate(function_tables, start=1):
            with self.subTest(function=index):
                self.assertEqual(table.cell(0, 0).text, f"FUNC-{index:02d}")
                self.assertEqual(
                    tuple(row.cells[0].text for row in table.rows[1:]),
                    expected_labels,
                )
                main_flow = table.cell(5, 1).text
                self.assertIn("1. ", main_flow)
                self.assertIn("5. ", main_flow)

    def test_tables_media_and_relationships_are_self_contained(self):
        self.assertGreaterEqual(len(self.document.tables), 20)
        for table_index, table in enumerate(self.document.tables, start=1):
            self.assertTrue(table.rows, f"table {table_index} has no rows")
            for row_index, row in enumerate(table.rows, start=1):
                for column_index, cell in enumerate(row.cells, start=1):
                    self.assertTrue(
                        cell.text.strip(),
                        f"empty table cell {table_index}:{row_index}:{column_index}",
                    )

    def test_portrait_ui_designs_keep_full_readable_page_height(self):
        portrait_shapes = [
            shape
            for shape in self.document.inline_shapes
            if shape.height / shape.width > 2
        ]
        self.assertEqual(len(portrait_shapes), 3)
        for shape in portrait_shapes:
            self.assertAlmostEqual(shape.height.cm, 20.5, places=1)
            self.assertGreaterEqual(shape.width.cm, 9.4)

        media = {
            name: data
            for name, data in self.parts.items()
            if name.startswith("word/media/")
        }
        self.assertGreaterEqual(len(media), 14)
        embedded_hashes = {sha256(data).hexdigest() for data in media.values()}
        generated_hashes = {
            sha256(path.read_bytes()).hexdigest()
            for path in self.ASSET_DIR.glob("*.png")
        }
        design_hashes = {
            sha256(path.read_bytes()).hexdigest() for path in self.DESIGN_IMAGES
        }
        self.assertEqual(len(generated_hashes), 11)
        self.assertLessEqual(generated_hashes | design_hashes, embedded_hashes)

        self.assertEqual(_relationship_violations(self.parts), [])

    def test_relationship_audit_is_attribute_order_independent_and_blocks_local_paths(self):
        targets = (
            "https://example.invalid/image.png",
            "file:///tmp/image.png",
            "/tmp/image.png",
            "C:/temp/image.png",
            r"\\server\share\image.png",
            "//server/share/image.png",
            "../../file.png",
        )
        relationships = "".join(
            (
                f'<Relationship Id="rId{index}" Target="{target}" '
                'TargetMode="External" Type="fixture"/>'
                if index <= 2
                else f'<Relationship Id="rId{index}" Target="{target}" Type="fixture"/>'
            )
            for index, target in enumerate(targets, start=1)
        )
        fixture = {
            "word/_rels/document.xml.rels": (
                f'<Relationships xmlns="{_RELATIONSHIPS_NS}">'
                f"{relationships}</Relationships>"
            ).encode("utf-8")
        }
        violations = _relationship_violations(fixture)
        violated_targets = {target for _part, target, _reason in violations}
        self.assertEqual(violated_targets, set(targets))

    def test_toc_page_fields_numbering_and_chapter_captions_are_native(self):
        settings = self.parts["word/settings.xml"].decode("utf-8")
        self.assertRegex(
            settings,
            r"<w:updateFields(?:\s+[^>]*)?\s+w:val=\"true\"",
        )

        field_xml = "\n".join(
            data.decode("utf-8")
            for name, data in self.parts.items()
            if name == "word/document.xml"
            or name.startswith("word/header")
            or name.startswith("word/footer")
        )
        self.assertIn('TOC \\o "1-3" \\h \\z \\u', field_xml)
        self.assertGreaterEqual(field_xml.count(">PAGE<"), 11)

        document_xml = self.parts["word/document.xml"].decode("utf-8")
        numbering_xml = self.parts["word/numbering.xml"].decode("utf-8")
        self.assertIn("<w:numPr>", document_xml)
        self.assertIn('w:val="bullet"', numbering_xml)
        self.assertIn('w:val="decimal"', numbering_xml)

        captions = [
            paragraph.text.strip()
            for paragraph in self.document.paragraphs
            if paragraph.style.name == "Caption"
        ]
        self.assertGreaterEqual(len(captions), 13)
        self.assertTrue(any(re.match(r"图5-\d+ ", text) for text in captions))
        self.assertTrue(any(re.match(r"图7-\d+ ", text) for text in captions))
        self.assertTrue(any(re.match(r"表10-\d+ ", text) for text in captions))

    def test_toc_cache_materialization_preserves_field_and_visible_page_leaders(self):
        from scripts.agent_a3_document.styles import (
            add_toc_field,
            configure_document,
        )

        try:
            from scripts.agent_a3_document.docx_builder import (
                TocEntry,
                materialize_toc_cache,
            )
        except ImportError as exc:
            self.fail(f"missing TOC cache materializer: {exc}")

        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "toc.docx"
            document = Document()
            configure_document(document)
            toc = document.add_paragraph()
            add_toc_field(toc)
            document.add_paragraph("第一章 项目概述", style="Heading 1")
            document.save(path)

            entries = (
                TocEntry(level=0, title="文档信息", page="I"),
                TocEntry(level=0, title="1.1 建设背景", page=7),
                TocEntry(level=1, title="1.1.1 事实边界", page=8),
            )
            try:
                materialized = materialize_toc_cache(path, entries)
            except TypeError as exc:
                self.fail(f"TOC materializer rejected Roman page labels: {exc}")
            self.assertEqual(materialized, path)

            with ZipFile(path) as package:
                document_xml = package.read("word/document.xml").decode("utf-8")

        self.assertIn('TOC \\o "1-3" \\h \\z \\u', document_xml)
        self.assertNotIn("目录将在 Word 中更新", document_xml)
        for entry in entries:
            self.assertIn(entry.title, document_xml)
            self.assertRegex(
                document_xml,
                rf"{re.escape(entry.title)}.*?<w:tab/>.*?<w:t>{entry.page}</w:t>",
            )
        self.assertEqual(document_xml.count('w:fldCharType="begin"'), 1)
        self.assertEqual(document_xml.count('w:fldCharType="separate"'), 1)
        self.assertEqual(document_xml.count('w:fldCharType="end"'), 1)
        self.assertEqual(document_xml.count('w:leader="dot"'), len(entries))

    def test_missing_or_unsafe_figure_assets_fail_clearly(self):
        from scripts.agent_a3_document.docx_builder import build_document

        for src, expected in (
            ("images/missing.png", "missing figure asset"),
            ("../outside.png", "repository-relative figure path"),
        ):
            with self.subTest(src=src), TemporaryDirectory() as tmp:
                root = Path(tmp)
                source_dir = root / "source"
                source_dir.mkdir()
                directive = (
                    f'<!-- FIGURE src="{src}" caption="验证图" '
                    'width_cm="15.5" -->\n'
                )
                _write_sources(source_dir, directive)
                output = root / "result.docx"
                with self.assertRaisesRegex(ValueError, expected):
                    build_document(
                        source_dir,
                        self.EVIDENCE_PATH,
                        self.ASSET_DIR,
                        output,
                    )
                self.assertFalse(output.exists())

    def test_generated_asset_symlink_cannot_escape_asset_directory(self):
        from scripts.agent_a3_document.docx_builder import build_document

        with TemporaryDirectory() as tmp:
            root = Path(tmp)
            asset_dir = root / "assets"
            asset_dir.mkdir()
            victim = "system-context.png"
            outside = root / "outside.png"
            copy2(self.ASSET_DIR / victim, outside)
            for source in self.ASSET_DIR.glob("*.png"):
                if source.name != victim:
                    copy2(source, asset_dir / source.name)
            try:
                (asset_dir / victim).symlink_to(outside)
            except (NotImplementedError, OSError) as exc:
                self.skipTest(f"symlinks unavailable: {exc}")

            output = root / "result.docx"
            with self.assertRaisesRegex(
                ValueError, r"generated asset escapes asset directory"
            ):
                build_document(
                    self.SOURCE_DIR,
                    self.EVIDENCE_PATH,
                    asset_dir,
                    output,
                )
            self.assertFalse(output.exists())

    def test_repository_design_symlink_cannot_escape_repository(self):
        from scripts.agent_a3_document.docx_builder import build_document

        with TemporaryDirectory() as repo_tmp, TemporaryDirectory() as outside_tmp:
            repository = Path(repo_tmp)
            (repository / ".git").mkdir()
            source_dir = repository / "docs/project-document/source"
            asset_dir = repository / "docs/project-document/assets/generated"
            design_dir = repository / "docs/designs/profile-radar"
            source_dir.mkdir(parents=True)
            asset_dir.mkdir(parents=True)
            design_dir.mkdir(parents=True)
            directive = (
                '<!-- FIGURE src="images/profile-radar-design.png" '
                'caption="界面设计稿" width_cm="15.5" -->\n'
            )
            _write_sources(source_dir, directive)
            for source in self.ASSET_DIR.glob("*.png"):
                copy2(source, asset_dir / source.name)

            outside = Path(outside_tmp) / "profile-radar-design.png"
            copy2(self.DESIGN_IMAGES[0], outside)
            link = design_dir / "profile-radar-design.png"
            try:
                link.symlink_to(outside)
            except (NotImplementedError, OSError) as exc:
                self.skipTest(f"symlinks unavailable: {exc}")

            output = repository / "result.docx"
            with self.assertRaisesRegex(
                ValueError, r"repository-relative figure path escapes repository"
            ):
                build_document(
                    source_dir,
                    self.EVIDENCE_PATH,
                    asset_dir,
                    output,
                )
            self.assertFalse(output.exists())


if __name__ == "__main__":
    unittest.main()
